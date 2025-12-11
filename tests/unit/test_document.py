"""
Unit tests for document generation modules
"""

from unittest.mock import patch

import pytest
from docx import Document as DocxDocument
from docx.shared import RGBColor

from solairus_intelligence.config.clients import ClientSector
from solairus_intelligence.core.document.content import ContentExtractor
from solairus_intelligence.core.document.generator import DocumentGenerator
from solairus_intelligence.core.document.sections import (
    EconomicIndicatorsBuilder,
    ExecutiveSummaryBuilder,
    HeaderBuilder,
    RegionalAssessmentBuilder,
    SectorSectionBuilder,
)
from solairus_intelligence.core.document.styles import (
    ERGO_COLORS,
    SPACING,
    ErgoStyles,
    FontConfig,
)
from solairus_intelligence.core.processors.base import IntelligenceItem, SectorIntelligence


class TestErgoColors:
    """Test Ergo brand colors"""

    def test_colors_defined(self):
        """Test colors are defined"""
        assert ERGO_COLORS is not None
        assert len(ERGO_COLORS) > 0

    def test_primary_colors_exist(self):
        """Test primary colors exist"""
        assert "primary_blue" in ERGO_COLORS
        assert "secondary_blue" in ERGO_COLORS
        assert "light_blue" in ERGO_COLORS

    def test_colors_are_rgb(self):
        """Test colors are RGBColor objects"""
        for name, color in ERGO_COLORS.items():
            assert isinstance(color, RGBColor)


class TestSpacing:
    """Test spacing constants"""

    def test_spacing_defined(self):
        """Test spacing is defined"""
        assert SPACING is not None
        assert len(SPACING) > 0

    def test_spacing_values(self):
        """Test spacing values exist"""
        assert "section_before" in SPACING
        assert "paragraph" in SPACING
        assert "bullet" in SPACING

    def test_spacing_are_integers(self):
        """Test spacing values are integers"""
        for name, value in SPACING.items():
            assert isinstance(value, int)


class TestFontConfig:
    """Test FontConfig dataclass"""

    def test_default_values(self):
        """Test default font configuration"""
        config = FontConfig()

        assert config.name == "Calibri"
        assert config.heading_size == 14
        assert config.subheading_size == 12
        assert config.body_size == 10

    def test_custom_values(self):
        """Test custom font configuration"""
        config = FontConfig(name="Arial", heading_size=16, body_size=11)

        assert config.name == "Arial"
        assert config.heading_size == 16
        assert config.body_size == 11


class TestErgoStyles:
    """Test ErgoStyles class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    def test_styles_initialization(self, styles):
        """Test styles initialize correctly"""
        assert styles is not None
        assert styles.colors is not None
        assert styles.spacing is not None
        assert styles.font_config is not None

    def test_styles_with_custom_font(self):
        """Test styles with custom font config"""
        font_config = FontConfig(name="Arial", body_size=11)
        styles = ErgoStyles(font_config=font_config)

        assert styles.font_config.name == "Arial"
        assert styles.font_config.body_size == 11

    def test_get_color(self, styles):
        """Test getting color by name"""
        color = styles.get_color("primary_blue")

        assert isinstance(color, RGBColor)

    def test_get_color_unknown(self, styles):
        """Test getting unknown color returns black"""
        color = styles.get_color("nonexistent_color")

        assert color == ERGO_COLORS["black"]

    def test_get_spacing(self, styles):
        """Test getting spacing by name"""
        spacing = styles.get_spacing("section_before")

        assert isinstance(spacing, int)
        assert spacing > 0

    def test_get_spacing_unknown(self, styles):
        """Test getting unknown spacing returns paragraph default"""
        spacing = styles.get_spacing("nonexistent_spacing")

        assert spacing == SPACING["paragraph"]

    def test_apply_to_document(self, styles):
        """Test applying styles to document"""
        doc = DocxDocument()

        styles.apply_to_document(doc)

        # Check that Normal style was modified
        normal_style = doc.styles["Normal"]
        assert normal_style.font.name == "Calibri"


class TestContentExtractor:
    """Test ContentExtractor class"""

    @pytest.fixture
    def extractor(self):
        return ContentExtractor()

    def test_extractor_initialization(self, extractor):
        """Test extractor initializes correctly"""
        assert extractor is not None

    def test_extractor_has_extract_method(self, extractor):
        """Test extractor has extract methods"""
        assert hasattr(extractor, "extract_analytical_insights")
        assert hasattr(extractor, "extract_theme")
        assert hasattr(extractor, "extract_value")

    def test_extract_theme_geopolitical(self, extractor):
        """Test extracting geopolitical theme"""
        result = extractor.extract_theme(
            "Military conflict escalating in the region", "Tensions affecting trade routes"
        )
        assert result == "Geopolitical Risk"

    def test_extract_theme_economic(self, extractor):
        """Test extracting economic theme"""
        result = extractor.extract_theme(
            "Inflation continues to rise", "Economic pressure mounting"
        )
        assert result == "Economic Pressure"

    def test_extract_theme_trade(self, extractor):
        """Test extracting trade theme"""
        result = extractor.extract_theme("New tariff announced on imports", "Trade policy impact")
        assert result == "Trade Policy"

    def test_extract_theme_default(self, extractor):
        """Test default theme"""
        result = extractor.extract_theme("General update", "No specific keywords")
        assert result == "Strategic Development"

    def test_extract_value_percentage(self, extractor):
        """Test extracting percentage value"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Rate increased by 5.5%",
            category="economic",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.extract_value(item)
        assert "5.5" in result or "%" in result

    def test_extract_value_dollar(self, extractor):
        """Test extracting dollar value"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Investment of $50 million",
            category="economic",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.extract_value(item)
        assert "50" in result

    def test_determine_trend_up(self, extractor):
        """Test determining upward trend"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Rate increased significantly",
            category="economic",
            relevance_score=0.8,
            so_what_statement="Growth observed",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.determine_trend(item)
        assert result == "↑"

    def test_determine_trend_down(self, extractor):
        """Test determining downward trend"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Rate decreased",
            category="economic",
            relevance_score=0.8,
            so_what_statement="Decline observed",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.determine_trend(item)
        assert result == "↓"

    def test_determine_trend_stable(self, extractor):
        """Test determining stable trend"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Rate stable",
            category="economic",
            relevance_score=0.8,
            so_what_statement="No change",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.determine_trend(item)
        assert result == "→"

    def test_strip_markdown(self, extractor):
        """Test stripping markdown"""
        text = "**Bold** and *italic* text"
        result = extractor.strip_markdown(text)
        assert "Bold" in result
        assert "**" not in result
        assert "*" not in result

    def test_strip_markdown_empty(self, extractor):
        """Test stripping markdown from empty string"""
        result = extractor.strip_markdown("")
        assert result == ""

    def test_parse_key_finding_with_colon(self, extractor):
        """Test parsing key finding with colon"""
        header, desc, bullets = extractor.parse_key_finding("Economic: Growth continues")
        assert header == "Economic"
        assert desc == "Growth continues"
        assert isinstance(bullets, list)

    def test_parse_key_finding_without_colon(self, extractor):
        """Test parsing key finding without colon"""
        header, desc, bullets = extractor.parse_key_finding("General observation")
        assert header == "Key Development"
        assert "General observation" in desc

    def test_craft_bottom_line_statement(self, extractor):
        """Test crafting bottom line statement"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Details here",
            category="economic",
            relevance_score=0.8,
            so_what_statement="This is the key insight",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.craft_bottom_line_statement(item)
        assert "key insight" in result

    def test_extract_indicator_name(self, extractor):
        """Test extracting indicator name"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="Data",
            category="inflation",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        result = extractor.extract_indicator_name(item)
        assert "Inflation" in result or "CPI" in result


class TestHeaderBuilder:
    """Test HeaderBuilder class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    @pytest.fixture
    def builder(self, styles):
        return HeaderBuilder(styles)

    def test_builder_initialization(self, builder):
        """Test header builder initializes correctly"""
        assert builder is not None
        assert builder.styles is not None

    def test_builder_with_logo_path(self, styles):
        """Test builder with logo path"""
        builder = HeaderBuilder(styles, logo_path="/path/to/logo.png")
        assert builder.logo_path == "/path/to/logo.png"

    def test_add_header(self, builder):
        """Test adding header"""
        doc = DocxDocument()
        builder.add_header(doc, "Test Report", "Test Subtitle")
        assert len(doc.paragraphs) >= 2


class TestExecutiveSummaryBuilder:
    """Test ExecutiveSummaryBuilder class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    @pytest.fixture
    def extractor(self):
        return ContentExtractor()

    @pytest.fixture
    def builder(self, styles, extractor):
        return ExecutiveSummaryBuilder(styles, extractor)

    def test_builder_initialization(self, builder):
        """Test executive summary builder initializes correctly"""
        assert builder is not None
        assert builder.styles is not None
        assert builder.content_extractor is not None

    def test_add_executive_summary_empty(self, builder):
        """Test adding executive summary with no items"""
        doc = DocxDocument()
        builder.add_executive_summary(doc, [])
        assert len(doc.paragraphs) > 0

    def test_add_executive_summary_with_items(self, builder):
        """Test adding executive summary with items"""
        doc = DocxDocument()
        items = [
            IntelligenceItem(
                raw_content="test",
                processed_content="Economic growth continues",
                category="economic",
                relevance_score=0.9,
                so_what_statement="Significant growth observed",
                affected_sectors=[ClientSector.GENERAL],
            )
        ]
        builder.add_executive_summary(doc, items)
        assert len(doc.paragraphs) > 0


class TestEconomicIndicatorsBuilder:
    """Test EconomicIndicatorsBuilder class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    @pytest.fixture
    def extractor(self):
        return ContentExtractor()

    @pytest.fixture
    def builder(self, styles, extractor):
        return EconomicIndicatorsBuilder(styles, extractor)

    def test_builder_initialization(self, builder):
        """Test economic indicators builder initializes correctly"""
        assert builder is not None
        assert builder.styles is not None

    def test_add_economic_indicators_empty(self, builder):
        """Test adding economic indicators with no items"""
        doc = DocxDocument()
        builder.add_economic_indicators_table(doc, [])
        # No table should be added
        assert len(doc.tables) == 0

    def test_add_economic_indicators_with_items(self, builder):
        """Test adding economic indicators with items"""
        doc = DocxDocument()
        items = [
            IntelligenceItem(
                raw_content="test",
                processed_content="CPI 3.2% increase",
                category="economic",
                relevance_score=0.8,
                so_what_statement="Inflation rising",
                affected_sectors=[ClientSector.GENERAL],
                source_type="fred",
            )
        ]
        builder.add_economic_indicators_table(doc, items)
        assert len(doc.tables) > 0


class TestRegionalAssessmentBuilder:
    """Test RegionalAssessmentBuilder class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    @pytest.fixture
    def extractor(self):
        return ContentExtractor()

    @pytest.fixture
    def builder(self, styles, extractor):
        return RegionalAssessmentBuilder(styles, extractor)

    def test_builder_initialization(self, builder):
        """Test regional assessment builder initializes correctly"""
        assert builder is not None
        assert builder.styles is not None

    def test_add_regional_assessment_empty(self, builder):
        """Test adding regional assessment with no items"""
        doc = DocxDocument()
        builder.add_regional_assessment(doc, [])
        # No content added
        assert len(doc.paragraphs) == 0

    def test_add_regional_assessment_with_items(self, builder):
        """Test adding regional assessment with items"""
        doc = DocxDocument()
        items = [
            IntelligenceItem(
                raw_content="Europe trade update",
                processed_content="European markets showing growth",
                category="trade",
                relevance_score=0.8,
                so_what_statement="EU impact",
                affected_sectors=[ClientSector.GENERAL],
            )
        ]
        builder.add_regional_assessment(doc, items)
        assert len(doc.paragraphs) > 0

    def test_detect_region_europe(self, builder):
        """Test detecting Europe region"""
        item = IntelligenceItem(
            raw_content="EU policy",
            processed_content="European Union update",
            category="policy",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        region = builder._detect_region(item)
        assert region == "Europe"

    def test_detect_region_asia(self, builder):
        """Test detecting Asia region"""
        item = IntelligenceItem(
            raw_content="China trade",
            processed_content="Chinese markets",
            category="trade",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        region = builder._detect_region(item)
        assert region == "Asia-Pacific"

    def test_detect_region_global(self, builder):
        """Test detecting Global region"""
        item = IntelligenceItem(
            raw_content="test",
            processed_content="General update",
            category="general",
            relevance_score=0.8,
            so_what_statement="Impact",
            affected_sectors=[ClientSector.GENERAL],
        )
        region = builder._detect_region(item)
        assert region == "Global"


class TestSectorSectionBuilder:
    """Test SectorSectionBuilder class"""

    @pytest.fixture
    def styles(self):
        return ErgoStyles()

    @pytest.fixture
    def extractor(self):
        return ContentExtractor()

    @pytest.fixture
    def builder(self, styles, extractor):
        return SectorSectionBuilder(styles, extractor)

    def test_builder_initialization(self, builder):
        """Test sector section builder initializes correctly"""
        assert builder is not None
        assert builder.styles is not None

    def test_add_sector_section(self, builder):
        """Test adding sector section"""
        doc = DocxDocument()
        intelligence = SectorIntelligence(
            sector=ClientSector.TECHNOLOGY,
            items=[
                IntelligenceItem(
                    raw_content="test",
                    processed_content="Tech sector update",
                    category="technology",
                    relevance_score=0.8,
                    so_what_statement="Impact on tech sector",
                    affected_sectors=[ClientSector.TECHNOLOGY],
                )
            ],
            summary="Technology sector summary",
        )
        builder.add_sector_section(doc, ClientSector.TECHNOLOGY, intelligence)
        assert len(doc.paragraphs) > 0


class TestDocumentGenerator:
    """Test DocumentGenerator class"""

    @pytest.fixture
    def generator(self):
        with patch.object(DocumentGenerator, "_init_ai_generator"):
            return DocumentGenerator()

    def test_generator_initialization(self, generator):
        """Test generator initializes correctly"""
        assert generator is not None
        assert generator.styles is not None
        assert generator.content_extractor is not None

    def test_generator_has_builders(self, generator):
        """Test generator has section builders"""
        assert generator.header_builder is not None
        assert generator.exec_summary_builder is not None
        assert generator.econ_indicators_builder is not None
        assert generator.regional_builder is not None
        assert generator.sector_builder is not None

    def test_generator_has_create_report(self, generator):
        """Test generator has create_report method"""
        assert hasattr(generator, "create_report")
        assert callable(generator.create_report)

    def test_generator_has_save_report(self, generator):
        """Test generator has save_report method"""
        assert hasattr(generator, "save_report")
        assert callable(generator.save_report)

    def test_create_report(self, generator):
        """Test creating a report"""
        items = [
            IntelligenceItem(
                raw_content="test",
                processed_content="Test content",
                category="economic",
                relevance_score=0.8,
                so_what_statement="Test impact",
                affected_sectors=[ClientSector.GENERAL],
            )
        ]
        sector_intel = {}

        doc = generator.create_report(items, sector_intel)
        # Check it's a document by verifying it has paragraphs
        assert doc is not None
        assert hasattr(doc, "paragraphs")

    def test_create_report_with_month(self, generator):
        """Test creating report with custom month"""
        doc = generator.create_report([], {}, report_month="January 2024")
        assert doc is not None
        assert hasattr(doc, "paragraphs")


class TestDocumentIntegration:
    """Integration tests for document generation"""

    def test_create_simple_document(self):
        """Test creating a simple document"""
        doc = DocxDocument()
        styles = ErgoStyles()
        styles.apply_to_document(doc)

        doc.add_heading("Test Report", level=0)
        doc.add_paragraph("Test content")

        assert len(doc.paragraphs) > 0

    def test_styles_chain_correctly(self):
        """Test that styles can be chained"""
        styles = ErgoStyles()
        doc = DocxDocument()

        styles.apply_to_document(doc)

        # Should be able to get colors and spacing after applying
        color = styles.get_color("primary_blue")
        spacing = styles.get_spacing("section_before")

        assert color is not None
        assert spacing > 0

    def test_full_document_workflow(self):
        """Test complete document generation workflow"""
        with patch.object(DocumentGenerator, "_init_ai_generator"):
            generator = DocumentGenerator()

        # Generator should have initialized properly
        assert generator.styles is not None
        assert generator.content_extractor is not None

    def test_header_and_summary_integration(self):
        """Test header and summary builders work together"""
        styles = ErgoStyles()
        extractor = ContentExtractor()
        doc = DocxDocument()

        header_builder = HeaderBuilder(styles)
        header_builder.add_header(doc, "Intelligence Report", "Monthly Analysis")

        summary_builder = ExecutiveSummaryBuilder(styles, extractor)
        summary_builder.add_executive_summary(doc, [])

        assert len(doc.paragraphs) > 0
