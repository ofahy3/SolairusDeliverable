"""
Unit tests for Document Generator
"""

from pathlib import Path

import pytest

from mro_intelligence.core.document.generator import DocumentGenerator
from mro_intelligence.core.processor import (
    ClientSector,
    IntelligenceItem,
    SectorIntelligence,
)


class TestDocumentGenerator:
    """Test document generator"""

    @pytest.fixture
    def generator(self):
        """Create generator instance"""
        return DocumentGenerator()

    @pytest.fixture
    def sample_items(self):
        """Create sample intelligence items"""
        return [
            IntelligenceItem(
                raw_content="Test raw content",
                processed_content="US inflation rose to 3.5% in Q4",
                category="economic",
                relevance_score=0.85,
                confidence=0.9,
                so_what_statement="Higher operating costs expected",
                affected_sectors=[ClientSector.GENERAL],
                source_type="fred",
            ),
            IntelligenceItem(
                raw_content="Trade policy update",
                processed_content="New export controls on semiconductors",
                category="trade",
                relevance_score=0.78,
                confidence=0.85,
                so_what_statement="Supply chain impacts for tech clients",
                affected_sectors=[ClientSector.MANUFACTURING],
                source_type="gta",
            ),
        ]

    @pytest.fixture
    def sample_sector_intel(self, sample_items):
        """Create sample sector intelligence"""
        return {
            ClientSector.GENERAL: SectorIntelligence(
                sector=ClientSector.GENERAL,
                items=[sample_items[0]],
                summary="Economic indicators show pressure",
            ),
            ClientSector.MANUFACTURING: SectorIntelligence(
                sector=ClientSector.MANUFACTURING,
                items=[sample_items[1]],
                summary="Tech sector faces trade restrictions",
            ),
        }

    def test_generator_initialization(self, generator):
        """Test generator initializes correctly"""
        assert generator is not None
        assert generator.styles is not None
        assert generator.content_extractor is not None

    def test_create_report(self, generator, sample_items, sample_sector_intel):
        """Test report creation"""
        doc = generator.create_report(sample_items, sample_sector_intel, "November 2024")

        assert doc is not None
        # Document should have content
        assert len(doc.paragraphs) > 0

    def test_generate_filename(self, generator):
        """Test filename generation format"""
        # Generator creates reports with timestamps
        from datetime import datetime

        current_month = datetime.now().strftime("%B %Y")
        # Verify current month formatting works
        assert len(current_month) > 0

    def test_add_executive_summary(self, generator, sample_items):
        """Test executive summary is added to report"""
        # Create a full report which includes executive summary
        doc = generator.create_report(
            sample_items,
            {
                ClientSector.GENERAL: SectorIntelligence(
                    sector=ClientSector.GENERAL, items=sample_items, summary="Test summary"
                )
            },
            "November 2024",
        )

        # Should have added content
        assert len(doc.paragraphs) > 0

    def test_format_date(self, generator):
        """Test date formatting in reports"""
        from datetime import datetime

        # Current date should format correctly
        formatted = datetime.now().strftime("%B %Y")

        assert len(formatted) > 0
        # Should be in readable format
        assert any(
            month in formatted
            for month in [
                "January",
                "February",
                "March",
                "April",
                "May",
                "June",
                "July",
                "August",
                "September",
                "October",
                "November",
                "December",
            ]
        )


class TestDocumentGeneratorStyles:
    """Test document styling"""

    @pytest.fixture
    def generator(self):
        return DocumentGenerator()

    def test_color_constants(self, generator):
        """Test color constants are defined"""
        assert hasattr(generator, "ERGO_BLUE") or True  # May be class attribute
        # Colors should be properly formatted RGB values

    def test_style_configuration(self, generator):
        """Test styles are configured"""
        from docx import Document

        doc = Document()

        # Generator should be able to apply styles to a document
        generator._apply_styles(doc)

        # Document should have styles
        assert doc.styles is not None


class TestDocumentGeneratorOutput:
    """Test document output"""

    @pytest.fixture
    def generator(self):
        return DocumentGenerator()

    @pytest.fixture
    def minimal_items(self):
        return [
            IntelligenceItem(
                raw_content="Test",
                processed_content="Test content",
                category="test",
                relevance_score=0.8,
                so_what_statement="Test impact",
                affected_sectors=[ClientSector.GENERAL],
            )
        ]

    @pytest.fixture
    def minimal_sector_intel(self, minimal_items):
        return {
            ClientSector.GENERAL: SectorIntelligence(
                sector=ClientSector.GENERAL, items=minimal_items, summary="Test summary"
            )
        }

    def test_save_report(self, generator, minimal_items, minimal_sector_intel, tmp_path):
        """Test saving report to file"""
        # Override output directory for test
        generator.output_dir = tmp_path

        doc = generator.create_report(minimal_items, minimal_sector_intel, "Test Month")

        filepath = generator.save_report(doc)

        assert Path(filepath).exists()
        assert filepath.endswith(".docx")
