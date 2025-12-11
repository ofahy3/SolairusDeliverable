"""
Main Document Generator for MRO Intelligence Reports.

Orchestrates the document generation process using focused modules.
Generates a ~3 page MRO Market Compendium with:
- Page 1: Executive Summary & Macro Outlook
- Page 2: Sector Demand Analysis
- Page 3: Risks & Opportunities
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches

from mro_intelligence.config.content_blocklist import check_content
from mro_intelligence.core.document.content import ContentExtractor
from mro_intelligence.core.document.sections import (
    EconomicIndicatorsBuilder,
    ExecutiveSummaryBuilder,
    HeaderBuilder,
    IntroductionBuilder,
    MacroOutlookBuilder,
    RisksOpportunitiesBuilder,
    SectorDemandBuilder,
)
from mro_intelligence.core.document.styles import ErgoStyles
from mro_intelligence.core.processor import (
    ClientSector,
    IntelligenceItem,
    SectorIntelligence,
)
from mro_intelligence.utils.config import ENV_CONFIG, get_output_dir

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """
    Generates professional DOCX reports optimized for Google Docs.

    Uses a modular architecture with separate components for:
    - Styles: Brand colors and formatting
    - Content: Data extraction and parsing
    - Sections: Individual report section builders

    Report Structure (3 pages):
    - Page 1: Executive Summary & Macro Outlook
    - Page 2: Sector Demand Analysis
    - Page 3: Risks & Opportunities
    """

    def __init__(self):
        # Initialize components
        self.styles = ErgoStyles()
        self.content_extractor = ContentExtractor()

        # Path to logo
        self.logo_path = str(
            Path(__file__).parent.parent.parent / "web" / "static" / "images" / "ergo_logo.png"
        )

        # Initialize AI generator if enabled
        self.ai_generator: Optional[Any] = None
        if ENV_CONFIG.ai_enabled:
            self._init_ai_generator()

        # Initialize section builders
        self.header_builder = HeaderBuilder(self.styles, self.logo_path)
        self.intro_builder = IntroductionBuilder(self.styles)
        self.exec_summary_builder = ExecutiveSummaryBuilder(
            self.styles, self.content_extractor, self.ai_generator
        )
        self.macro_outlook_builder = MacroOutlookBuilder(self.styles, self.content_extractor)
        self.econ_indicators_builder = EconomicIndicatorsBuilder(
            self.styles, self.content_extractor
        )
        self.sector_demand_builder = SectorDemandBuilder(self.styles, self.content_extractor)
        self.risks_opps_builder = RisksOpportunitiesBuilder(self.styles, self.content_extractor)

        self.output_dir = get_output_dir()

    def _init_ai_generator(self) -> None:
        """Initialize AI generator for enhanced summaries"""
        try:
            from mro_intelligence.ai.generator import AIConfig, SecureAIGenerator

            api_key = ENV_CONFIG.anthropic_api_key
            if not api_key:
                logger.info("AI generation disabled: no API key configured")
                return

            config = AIConfig(api_key=api_key, model=ENV_CONFIG.ai_model, enabled=True)
            self.ai_generator = SecureAIGenerator(config)
            logger.info("AI-enhanced Executive Summary enabled")
        except Exception as e:
            logger.warning(f"AI generator initialization failed: {e}")

    def create_report(
        self,
        mro_items: List[IntelligenceItem],
        sector_intelligence: Dict[ClientSector, SectorIntelligence],
        report_month: Optional[str] = None,
    ) -> Document:
        """
        Create a complete intelligence report.

        Args:
            mro_items: List of intelligence items
            sector_intelligence: Dictionary of sector-specific intelligence
            report_month: Month for the report (e.g., "December 2024")

        Returns:
            Generated Document object
        """
        doc = Document()
        self._setup_document(doc)

        # Use current month if not specified
        if not report_month:
            report_month = datetime.now().strftime("%B %Y")

        # Build Page 1: Executive Summary & Macro Outlook
        self._create_page_1(doc, mro_items, report_month)

        # Add page break
        doc.add_page_break()

        # Build Page 2: Sector Demand Analysis
        self._create_page_2(doc, sector_intelligence, report_month)

        # Add page break
        doc.add_page_break()

        # Build Page 3: Risks & Opportunities
        self._create_page_3(doc, mro_items, sector_intelligence, report_month)

        return doc

    def _setup_document(self, doc: Document) -> None:
        """Set up document properties and styles"""
        # Document properties
        core_props = doc.core_properties
        core_props.author = "Ergo Intelligence"
        core_props.title = "MRO Market Intelligence Compendium"

        # Apply styles
        self.styles.apply_to_document(doc)

        # Set margins
        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

    def _create_page_1(self, doc: Document, items: List[IntelligenceItem], month: str) -> None:
        """
        Create Page 1: Executive Summary & Macro Outlook

        Sections:
        - Header with logo
        - Introduction explaining the analysis
        - Executive Summary: Top 4-5 key findings for MRO market
        - US Economic Outlook: GDP, industrial production, manufacturing trends
        - Key Economic Indicators table with MRO relevance
        """
        # Header
        self.header_builder.add_header(doc, "MRO Market Intelligence", month)

        # Introduction
        self.intro_builder.add_introduction(doc)

        # Executive Summary
        self.exec_summary_builder.add_executive_summary(doc, items)

        # US Economic Outlook
        self.macro_outlook_builder.add_macro_outlook(doc, items)

        # Key Economic Indicators table
        self.econ_indicators_builder.add_economic_indicators_table(doc, items)

    def _create_page_2(
        self, doc: Document, sector_intelligence: Dict[ClientSector, SectorIntelligence], month: str
    ) -> None:
        """
        Create Page 2: Sector Demand Analysis

        Sections for each sector:
        - Manufacturing: Production trends, capex, automation
        - Construction: Building activity, infrastructure spending
        - Energy: Oil/gas activity, utility maintenance
        - "So What for Grainger" callout boxes
        """
        # Header
        self.header_builder.add_header(doc, "Sector Demand Analysis", month)

        # Grainger's customer segments in order of importance
        priority_sectors = [
            ClientSector.MANUFACTURING,
            ClientSector.GOVERNMENT,
            ClientSector.COMMERCIAL_FACILITIES,
            ClientSector.CONTRACTORS,
        ]

        # Add priority sectors first
        for sector in priority_sectors:
            if sector in sector_intelligence and sector_intelligence[sector].items:
                self.sector_demand_builder.add_sector_demand_analysis(
                    doc, sector, sector_intelligence[sector]
                )

        # Add remaining sectors with content
        for sector, intelligence in sector_intelligence.items():
            if sector not in priority_sectors and intelligence.items:
                self.sector_demand_builder.add_sector_demand_analysis(doc, sector, intelligence)

    def _create_page_3(
        self,
        doc: Document,
        items: List[IntelligenceItem],
        sector_intelligence: Dict[ClientSector, SectorIntelligence],
        month: str,
    ) -> None:
        """
        Create Page 3: Risks & Opportunities

        Sections:
        - Trade Policy Impacts: Tariffs, supply chain shifts, reshoring
        - Regional Considerations: US domestic (USMCA region) focus
        - 90-Day Outlook: Near-term actionable intelligence
        - Recommended Actions: Concrete next steps
        """
        # Header
        self.header_builder.add_header(doc, "Risks & Opportunities", month)

        # Trade Policy Impacts
        self.risks_opps_builder.add_trade_policy_section(doc, items)

        # Regional Considerations
        self.risks_opps_builder.add_regional_section(doc, items)

        # 90-Day Outlook
        self.risks_opps_builder.add_outlook_section(doc, items)

        # Recommended Actions
        self.risks_opps_builder.add_recommended_actions(doc, items, sector_intelligence)

    def validate_output(self, content: str) -> None:
        """
        Validate that output contains no blocked content.
        Raises ValueError if contamination detected.

        Args:
            content: Text content to validate

        Raises:
            ValueError: If blocked content is found
        """
        violations = check_content(content)
        if violations:
            raise ValueError(
                f"CONTAMINATION DETECTED - Report generation blocked:\n"
                + "\n".join(violations)
            )

    def _extract_document_text(self, doc: Document) -> str:
        """Extract all text from a document for validation"""
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return " ".join(text_parts)

    def save_report(self, doc: Document, filename: Optional[str] = None) -> str:
        """
        Save the report to a file.

        Args:
            doc: Document to save
            filename: Optional custom filename

        Returns:
            Path to saved file

        Raises:
            ValueError: If document contains blocked content
        """
        # Validate document content before saving
        doc_text = self._extract_document_text(doc)
        self.validate_output(doc_text)

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            month = datetime.now().strftime("%B_%Y")
            filename = f"MRO_Intelligence_Report_{month}_{timestamp}.docx"

        # Ensure output directory exists
        self.output_dir.mkdir(parents=True, exist_ok=True)

        filepath = self.output_dir / filename
        doc.save(str(filepath))

        logger.info(f"Report saved: {filepath}")
        return str(filepath)

    # Backwards compatibility methods
    def _apply_styles(self, doc: Document) -> None:
        """Backwards compatible style application"""
        self.styles.apply_to_document(doc)

    def _setup_document_properties(self, doc: Document) -> None:
        """Backwards compatible setup"""
        self._setup_document(doc)
