"""
Section builders for MRO Intelligence Reports.

Each class is responsible for building a specific section of the report.
Restructured for 3-page MRO Market Compendium format:
- Page 1: Executive Summary & Macro Outlook
- Page 2: Sector Demand Analysis
- Page 3: Risks & Opportunities
"""

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from mro_intelligence.core.document.content import ContentExtractor
from mro_intelligence.core.document.styles import SPACING, ErgoStyles
from mro_intelligence.core.processor import (
    ClientSector,
    IntelligenceItem,
    SectorIntelligence,
)

logger = logging.getLogger(__name__)


class HeaderBuilder:
    """Builds document headers"""

    def __init__(self, styles: ErgoStyles, logo_path: Optional[str] = None):
        self.styles = styles
        self.logo_path = logo_path

    def add_header(self, doc: Document, title: str, subtitle: str) -> None:
        """Add branded header to document"""
        # Add logo if available
        if self.logo_path:
            try:
                from pathlib import Path

                if Path(self.logo_path).exists():
                    doc.add_picture(str(self.logo_path), width=Inches(1.5))
            except (FileNotFoundError, ValueError, OSError) as e:
                logging.debug(f"Could not add logo: {e}")  # Continue without logo

        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = self.styles.get_color("primary_blue")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = self.styles.get_color("dark_gray")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


class IntroductionBuilder:
    """Builds the introductory section explaining the analysis"""

    def __init__(self, styles: ErgoStyles):
        self.styles = styles

    def add_introduction(self, doc: Document) -> None:
        """Add introduction section explaining the analysis purpose"""
        # Introduction text
        intro_para = doc.add_paragraph()
        intro_text = (
            "This analysis translates Ergo Flashpoints Forum intelligence from the past "
            "3 months into implications for the MRO (Maintenance, Repair & Operations) market, "
            "tailored specifically for Grainger's business priorities."
        )
        intro_run = intro_para.add_run(intro_text)
        intro_run.font.size = Pt(10)
        intro_run.font.italic = True
        intro_run.font.color.rgb = self.styles.get_color("dark_gray")
        intro_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro_para.space_after = Pt(SPACING["section_after"])


class ExecutiveSummaryBuilder:
    """Builds the Executive Summary section with top key findings"""

    def __init__(
        self,
        styles: ErgoStyles,
        content_extractor: ContentExtractor,
        ai_generator: Optional[Any] = None,
    ):
        self.styles = styles
        self.content_extractor = content_extractor
        self.ai_generator = ai_generator

    def add_executive_summary(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add executive summary section with top 4-5 key findings for MRO"""
        self._add_section_heading(doc, "Executive Summary")

        if not items:
            doc.add_paragraph("No intelligence items available for analysis.")
            return

        # Intro sentence
        intro_para = doc.add_paragraph()
        intro_run = intro_para.add_run(
            "Top findings with direct implications for MRO market demand:"
        )
        intro_run.font.size = Pt(10)
        intro_run.font.italic = True
        intro_para.space_after = Pt(4)

        # Extract insights and create key findings list
        insights = self.content_extractor.extract_analytical_insights(items)

        # Combine bottom line and key findings for executive summary
        all_findings = []
        all_findings.extend(insights.get("bottom_line", []))
        all_findings.extend(insights.get("key_findings", []))

        # Display top 5 findings
        for i, finding in enumerate(all_findings[:5], 1):
            header, desc, _ = self.content_extractor.parse_key_finding(finding)

            finding_para = doc.add_paragraph()
            num_run = finding_para.add_run(f"{i}. ")
            num_run.font.bold = True
            num_run.font.size = Pt(10)
            num_run.font.color.rgb = self.styles.get_color("primary_blue")

            header_run = finding_para.add_run(f"{header}: ")
            header_run.font.bold = True
            header_run.font.size = Pt(10)

            desc_run = finding_para.add_run(desc)
            desc_run.font.size = Pt(10)
            finding_para.space_after = Pt(4)

    def _add_section_heading(self, doc: Document, heading: str) -> None:
        """Add a section heading"""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")
        para.space_before = Pt(SPACING["section_before"])
        para.space_after = Pt(SPACING["section_after"])


class MacroOutlookBuilder:
    """Builds the US Economic Outlook section"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_macro_outlook(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add US Economic Outlook section with GDP, industrial production, manufacturing"""
        self._add_section_heading(doc, "US Economic Outlook")

        # Filter for macro economic items
        macro_keywords = ["gdp", "industrial", "manufacturing", "production", "economic"]
        macro_items = [
            item for item in items
            if any(kw in (item.processed_content + item.category).lower() for kw in macro_keywords)
        ]

        if not macro_items:
            para = doc.add_paragraph()
            run = para.add_run(
                "Economic outlook remains stable with moderate industrial activity. "
                "Monitor FRED indicators for real-time updates on manufacturing and production trends."
            )
            run.font.size = Pt(10)
            return

        # Summary paragraph
        summary_para = doc.add_paragraph()
        summary_items = sorted(macro_items, key=lambda x: x.relevance_score, reverse=True)[:3]
        summary_text = " ".join([
            item.so_what_statement or item.processed_content[:150]
            for item in summary_items
        ])
        summary_run = summary_para.add_run(summary_text[:400] + "..." if len(summary_text) > 400 else summary_text)
        summary_run.font.size = Pt(10)

    def _add_section_heading(self, doc: Document, heading: str) -> None:
        """Add a section heading"""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("secondary_blue")
        para.space_before = Pt(SPACING["subsection_before"])
        para.space_after = Pt(SPACING["subsection_after"])


class EconomicIndicatorsBuilder:
    """Builds the Economic Indicators table with MRO relevance"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_economic_indicators_table(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add economic indicators table showing FRED data with MRO relevance"""
        # Filter for economic items
        econ_items = [
            item
            for item in items
            if item.source_type == "fred" or "economic" in item.category.lower()
        ]

        if not econ_items:
            return

        # Add section heading
        para = doc.add_paragraph()
        run = para.add_run("Key Economic Indicators")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("secondary_blue")
        para.space_before = Pt(SPACING["subsection_before"])
        para.space_after = Pt(SPACING["subsection_after"])

        # Create table with 4 columns
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        widths = [Inches(1.5), Inches(1.0), Inches(0.8), Inches(3.0)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # Header row with blue background
        headers = ["Indicator", "Value", "Trend", "MRO Relevance"]
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Style header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Blue background
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="1B2946" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows - limit to 8 indicators for Page 1
        for item in econ_items[:8]:
            row = table.add_row()
            row.cells[0].text = self._get_indicator_name(item)
            row.cells[1].text = self.content_extractor.extract_value(item)
            row.cells[2].text = self.content_extractor.determine_trend(item)
            row.cells[3].text = self._get_mro_relevance(item)

            # Style cells
            for idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx < 3 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.size = Pt(9)

    def _get_indicator_name(self, item: IntelligenceItem) -> str:
        """Get indicator name from FRED item"""
        # Check for FRED series name
        if hasattr(item, 'fred_series_id') and item.fred_series_id:
            series_names = {
                "INDPRO": "Industrial Production",
                "IPMAN": "Manufacturing Production",
                "DGORDER": "Durable Goods Orders",
                "HOUST": "Housing Starts",
                "PERMIT": "Building Permits",
                "TLRESCONS": "Residential Construction",
                "TLNRESCONS": "Nonresidential Construction",
                "UNRATE": "Unemployment Rate",
                "FEDFUNDS": "Federal Funds Rate",
                "PCEPILFE": "Core PCE Inflation",
                "PPIACO": "Producer Price Index",
                "WPU101": "Iron & Steel PPI",
                "DCOILWTICO": "Crude Oil (WTI)",
                "MANEMP": "Manufacturing Employment",
                "USCONS": "Construction Employment",
                "T10Y2Y": "Yield Curve Spread",
            }
            return series_names.get(item.fred_series_id, item.fred_series_id)

        return self.content_extractor.extract_indicator_name(item)

    def _get_mro_relevance(self, item: IntelligenceItem) -> str:
        """Get MRO market relevance statement"""
        if item.so_what_statement:
            # Truncate to fit table cell
            relevance = item.so_what_statement[:80]
            if len(item.so_what_statement) > 80:
                relevance = relevance.rsplit(" ", 1)[0] + "..."
            return relevance
        return "Monitor for MRO demand impact"


class SectorDemandBuilder:
    """Builds sector-specific demand analysis sections for Page 2"""

    # Sector display configuration - Grainger's customer segments
    SECTOR_CONFIG = {
        ClientSector.MANUFACTURING: {
            "title": "Manufacturing Customers",
            "focus_areas": ["production trends", "capex investment", "automation adoption"],
            "icon": "M",
        },
        ClientSector.GOVERNMENT: {
            "title": "Government & Defense",
            "focus_areas": ["federal spending", "defense budget", "infrastructure bills"],
            "icon": "G",
        },
        ClientSector.COMMERCIAL_FACILITIES: {
            "title": "Commercial Facilities",
            "focus_areas": ["occupancy trends", "building maintenance", "healthcare expansion"],
            "icon": "F",
        },
        ClientSector.CONTRACTORS: {
            "title": "Contractors",
            "focus_areas": ["construction activity", "housing starts", "permit trends"],
            "icon": "C",
        },
    }

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_sector_demand_analysis(
        self,
        doc: Document,
        sector: ClientSector,
        intelligence: SectorIntelligence,
    ) -> None:
        """Add sector demand analysis with So What for Grainger callout"""
        config = self.SECTOR_CONFIG.get(sector, {
            "title": f"{sector.value.title()} Sector",
            "focus_areas": [],
            "icon": "S",
        })

        # Sector heading
        self._add_sector_heading(doc, config["title"])

        # Focus areas subheading
        if config["focus_areas"]:
            focus_para = doc.add_paragraph()
            focus_run = focus_para.add_run(f"Focus: {', '.join(config['focus_areas'])}")
            focus_run.font.size = Pt(9)
            focus_run.font.italic = True
            focus_run.font.color.rgb = self.styles.get_color("dark_gray")
            focus_para.space_after = Pt(4)

        # Sector summary
        if intelligence.summary:
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(intelligence.summary)
            summary_run.font.size = Pt(10)
            summary_para.space_after = Pt(6)

        # Key items (limit to 3 per sector for page space)
        for item in intelligence.items[:3]:
            self._add_sector_item(doc, item)

        # "So What for Grainger" callout box
        self._add_grainger_callout(doc, sector, intelligence)

    def _add_sector_heading(self, doc: Document, heading: str) -> None:
        """Add sector heading"""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")
        para.space_before = Pt(SPACING["subsection_before"])
        para.space_after = Pt(4)

    def _add_sector_item(self, doc: Document, item: IntelligenceItem) -> None:
        """Add a sector intelligence item"""
        item_para = doc.add_paragraph()
        bullet_run = item_para.add_run("  ")
        bullet_run.font.size = Pt(10)

        content = item.so_what_statement or item.processed_content[:180]
        content_run = item_para.add_run(content)
        content_run.font.size = Pt(10)
        item_para.space_after = Pt(2)

    def _add_grainger_callout(
        self,
        doc: Document,
        sector: ClientSector,
        intelligence: SectorIntelligence,
    ) -> None:
        """Add 'So What for Grainger' callout box"""
        # Create callout table (1 row, 1 cell) for visual box
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        cell.width = Inches(6.5)

        # Light blue background
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="E5EAEF" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

        # Callout header
        header_para = cell.paragraphs[0]
        header_run = header_para.add_run("So What for Grainger")
        header_run.font.bold = True
        header_run.font.size = Pt(10)
        header_run.font.color.rgb = self.styles.get_color("primary_blue")

        # Generate Grainger-specific implication
        implication = self._generate_grainger_implication(sector, intelligence)
        impl_para = cell.add_paragraph()
        impl_run = impl_para.add_run(implication)
        impl_run.font.size = Pt(9)

        # Add spacing after table
        spacer = doc.add_paragraph()
        spacer.space_after = Pt(SPACING["paragraph"])

    def _generate_grainger_implication(
        self,
        sector: ClientSector,
        intelligence: SectorIntelligence,
    ) -> str:
        """Generate Grainger-specific business implication"""
        # Sector-specific implications for Grainger's customer segments
        implications = {
            ClientSector.MANUFACTURING: (
                "Manufacturing activity directly drives MRO demand for industrial supplies, "
                "safety equipment, and maintenance products. Production volumes correlate "
                "with consumables spend - monitor PMI for demand signals."
            ),
            ClientSector.GOVERNMENT: (
                "Federal and defense spending impacts Grainger's $2B+ government segment. "
                "Track appropriations bills, defense budget, and infrastructure spending "
                "for procurement planning and GSA contract positioning."
            ),
            ClientSector.COMMERCIAL_FACILITIES: (
                "Commercial occupancy and building activity drives HVAC, janitorial, and "
                "maintenance supply demand. Monitor return-to-office trends and healthcare "
                "facility expansion for demand signals."
            ),
            ClientSector.CONTRACTORS: (
                "Construction activity translates to demand for tools, fasteners, electrical, "
                "and plumbing supplies. Track housing starts and permits for forward-looking "
                "demand signals in contractor-focused categories."
            ),
        }

        base_implication = implications.get(
            sector,
            "Monitor sector developments for potential impact on product demand and customer needs."
        )

        # Enhance with actual intelligence if available
        if intelligence.items and intelligence.items[0].action_items:
            actions = intelligence.items[0].action_items[:2]
            action_text = " Action: " + "; ".join(actions) + "."
            return base_implication + action_text

        return base_implication


class RisksOpportunitiesBuilder:
    """Builds Page 3: Risks & Opportunities section"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_trade_policy_section(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add Trade Policy Impacts section"""
        self._add_section_heading(doc, "Trade Policy Impacts")

        # Filter for trade-related items
        trade_keywords = ["tariff", "trade", "import", "export", "sanction", "reshoring", "supply chain"]
        trade_items = [
            item for item in items
            if any(kw in (item.processed_content + item.category + (item.so_what_statement or "")).lower()
                   for kw in trade_keywords)
            or item.source_type == "gta"
        ]

        if not trade_items:
            para = doc.add_paragraph()
            run = para.add_run(
                "Trade environment remains stable. Continue monitoring Global Trade Alert "
                "for tariff changes and supply chain disruptions affecting industrial products."
            )
            run.font.size = Pt(10)
            return

        # Top trade impacts
        sorted_items = sorted(trade_items, key=lambda x: x.relevance_score, reverse=True)
        for item in sorted_items[:4]:
            item_para = doc.add_paragraph()
            bullet_run = item_para.add_run("  ")
            bullet_run.font.size = Pt(10)

            content = item.so_what_statement or item.processed_content[:200]
            content_run = item_para.add_run(content)
            content_run.font.size = Pt(10)
            item_para.space_after = Pt(4)

    def add_regional_section(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add Regional Considerations section - US domestic (USMCA) focus"""
        self._add_section_heading(doc, "Regional Considerations")

        # Subsection intro
        intro_para = doc.add_paragraph()
        intro_run = intro_para.add_run("Focus: US Domestic & USMCA Region")
        intro_run.font.size = Pt(10)
        intro_run.font.italic = True
        intro_run.font.color.rgb = self.styles.get_color("dark_gray")
        intro_para.space_after = Pt(4)

        # Filter for US/USMCA items
        regional_keywords = ["us", "usa", "america", "canada", "mexico", "usmca", "domestic", "north america"]
        regional_items = [
            item for item in items
            if any(kw in (item.processed_content + (item.so_what_statement or "")).lower()
                   for kw in regional_keywords)
        ]

        if not regional_items:
            para = doc.add_paragraph()
            run = para.add_run(
                "US and USMCA region conditions remain the primary focus for Grainger operations. "
                "Monitor nearshoring trends and cross-border trade developments."
            )
            run.font.size = Pt(10)
            return

        # Regional highlights
        sorted_items = sorted(regional_items, key=lambda x: x.relevance_score, reverse=True)
        for item in sorted_items[:3]:
            item_para = doc.add_paragraph()
            bullet_run = item_para.add_run("  ")
            bullet_run.font.size = Pt(10)

            content = item.so_what_statement or item.processed_content[:180]
            content_run = item_para.add_run(content)
            content_run.font.size = Pt(10)
            item_para.space_after = Pt(4)

    def add_outlook_section(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add 90-Day Outlook section"""
        self._add_section_heading(doc, "90-Day Outlook")

        # Generate forward-looking insights
        insights = self.content_extractor.extract_analytical_insights(items)
        watch_factors = insights.get("watch_factors", [])

        if not watch_factors:
            para = doc.add_paragraph()
            run = para.add_run(
                "Near-term outlook: Monitor economic indicators for demand signals. "
                "Industrial production and construction activity remain key drivers for MRO demand. "
                "Watch for Fed policy decisions and their impact on equipment financing."
            )
            run.font.size = Pt(10)
            return

        # Display watch factors as outlook items
        for i, factor in enumerate(watch_factors[:4], 1):
            title, desc, _ = self.content_extractor.parse_watch_factor(factor)

            factor_para = doc.add_paragraph()
            num_run = factor_para.add_run(f"{i}. ")
            num_run.font.bold = True
            num_run.font.size = Pt(10)
            num_run.font.color.rgb = self.styles.get_color("orange")

            title_run = factor_para.add_run(f"{title}: ")
            title_run.font.bold = True
            title_run.font.size = Pt(10)

            desc_run = factor_para.add_run(desc)
            desc_run.font.size = Pt(10)
            factor_para.space_after = Pt(4)

    def add_recommended_actions(
        self,
        doc: Document,
        items: List[IntelligenceItem],
        sector_intelligence: Dict[ClientSector, SectorIntelligence],
    ) -> None:
        """Add Recommended Actions section"""
        self._add_section_heading(doc, "Recommended Actions")

        # Collect action items from all sources
        all_actions = []

        # From high-relevance items
        sorted_items = sorted(items, key=lambda x: x.relevance_score, reverse=True)
        for item in sorted_items[:10]:
            if item.action_items:
                all_actions.extend(item.action_items[:2])

        # From sector intelligence
        for sector, intel in sector_intelligence.items():
            for item in intel.items[:3]:
                if item.action_items:
                    all_actions.extend(item.action_items[:1])

        # Deduplicate and limit
        unique_actions = list(dict.fromkeys(all_actions))[:6]

        if not unique_actions:
            unique_actions = [
                "Review inventory levels for high-demand MRO categories",
                "Monitor FRED economic indicators for demand signals",
                "Track Global Trade Alert for supply chain disruptions",
                "Engage with manufacturing customers on capex plans",
                "Assess pricing strategy for commodity-sensitive products",
            ]

        # Display as numbered list
        for i, action in enumerate(unique_actions, 1):
            action_para = doc.add_paragraph()
            num_run = action_para.add_run(f"{i}. ")
            num_run.font.bold = True
            num_run.font.size = Pt(10)
            num_run.font.color.rgb = self.styles.get_color("teal")

            action_run = action_para.add_run(action)
            action_run.font.size = Pt(10)
            action_para.space_after = Pt(4)

    def _add_section_heading(self, doc: Document, heading: str) -> None:
        """Add a section heading"""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("secondary_blue")
        para.space_before = Pt(SPACING["subsection_before"])
        para.space_after = Pt(SPACING["subsection_after"])
