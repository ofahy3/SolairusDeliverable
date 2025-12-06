"""
Document Generator for Solairus Intelligence Reports
Creates professional, Google Docs-compatible DOCX files with elegant formatting
"""

import asyncio
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from solairus_intelligence.core.processor import ClientSector, IntelligenceItem, SectorIntelligence
from solairus_intelligence.utils.config import ENV_CONFIG, get_output_dir


class DocumentGenerator:
    """
    Generates professional DOCX reports optimized for Google Docs
    With optional AI-enhanced Executive Summary generation
    """

    def __init__(self):
        # Ergo brand colors (full palette)
        self.ergo_colors = {
            # Primary colors
            "primary_blue": RGBColor(27, 41, 70),  # #1B2946 (colorP1)
            "secondary_blue": RGBColor(2, 113, 195),  # #0271C3 (colorP2)
            "light_blue": RGBColor(148, 176, 201),  # #94B0C9 (colorP3)
            # Secondary colors
            "orange": RGBColor(232, 116, 73),  # #E87449 (colorS1)
            "teal": RGBColor(0, 221, 204),  # #00DDCC (colorS2)
            "dark_navy": RGBColor(19, 31, 51),  # #131F33 (colorS3)
            # Tertiary colors
            "deep_orange": RGBColor(196, 97, 60),  # #C4613C (colorT1)
            "light_gray": RGBColor(229, 234, 239),  # #E5EAEF (colorT2)
            "dark_teal": RGBColor(1, 185, 171),  # #01B9AB (colorT3)
            "dark_blue": RGBColor(1, 93, 162),  # #015DA2 (colorT4)
            # Base colors
            "dark_gray": RGBColor(100, 100, 100),
            "black": RGBColor(0, 0, 0),
            "white": RGBColor(255, 255, 255),
        }

        # Standardized spacing constants (in points) for consistent document layout
        self.spacing = {
            "section_before": 18,  # Space before major section headings
            "section_after": 8,  # Space after section headings
            "subsection_before": 12,  # Space before subsection headings
            "subsection_after": 6,  # Space after subsection headings
            "paragraph": 8,  # Standard paragraph spacing
            "bullet": 4,  # Space after bullet items
            "table_after": 12,  # Space after tables
            "header_after": 6,  # Space after header elements
        }

        # Path to logo
        self.logo_path = (
            Path(__file__).parent.parent / "web" / "static" / "images" / "ergo_logo.png"
        )
        # Initialize AI generator if enabled
        self.ai_generator = None
        if ENV_CONFIG.ai_enabled:
            try:
                from solairus_intelligence.ai.generator import AIConfig, SecureAIGenerator

                config = AIConfig(
                    api_key=ENV_CONFIG.anthropic_api_key, model=ENV_CONFIG.ai_model, enabled=True
                )
                self.ai_generator = SecureAIGenerator(config)
                import logging

                logging.getLogger(__name__).info("✓ AI-enhanced Executive Summary enabled")
            except Exception as e:
                import logging

                logging.getLogger(__name__).warning(f"AI generator initialization failed: {e}")

    def create_report(
        self,
        solairus_items: List[IntelligenceItem],
        sector_intelligence: Dict[ClientSector, SectorIntelligence],
        report_month: Optional[str] = None,
    ) -> Document:
        """
        Create a complete two-page intelligence report
        """
        doc = Document()
        self._setup_document_properties(doc)
        self._apply_styles(doc)

        # Use current month if not specified
        if not report_month:
            report_month = datetime.now().strftime("%B %Y")

        # Page 1: Solairus Business Intelligence
        self._create_page_1(doc, solairus_items, report_month)

        # Client Sector Intelligence section removed per user request

        return doc

    def _setup_document_properties(self, doc: Document):
        """Set up document properties for optimal Google Docs compatibility"""
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _apply_styles(self, doc: Document):
        """Apply professional styling to the document"""
        # Modify default paragraph font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"  # Google Docs compatible
        font.size = Pt(11)
        font.color.rgb = self.ergo_colors["black"]

        # Create custom heading styles
        self._create_custom_heading_style(doc, "CustomHeading1", 16, True)
        self._create_custom_heading_style(doc, "CustomHeading2", 14, True)
        self._create_custom_heading_style(doc, "CustomHeading3", 12, False)

    def _create_custom_heading_style(self, doc: Document, name: str, size: int, bold: bool):
        """Create a custom heading style with standardized spacing"""
        styles = doc.styles
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            font = style.font
            font.size = Pt(size)
            font.bold = bold
            font.color.rgb = self.ergo_colors["primary_blue"]
            paragraph_format = style.paragraph_format
            paragraph_format.space_before = Pt(self.spacing["section_before"])
            paragraph_format.space_after = Pt(self.spacing["section_after"])
    def _create_page_1(self, doc: Document, items: List[IntelligenceItem], month: str):
        """Create Page 1: Solairus Business Intelligence"""

        # Header
        self._add_header(doc, "FLASHPOINTS FORUM INTELLIGENCE BRIEF", month)

        # Executive Summary
        self._add_section_heading(doc, "EXECUTIVE SUMMARY")
        self._add_executive_summary(doc, items)

        # Economic Indicators - as a table
        self._add_section_heading(doc, "ECONOMIC INDICATORS FOR BUSINESS AVIATION")
        econ_items = [
            item
            for item in items
            if "economic" in item.category
            or "financial" in item.category
            or "market" in item.category
        ]
        self._add_economic_indicators_table(doc, econ_items)

    def _create_page_2(
        self, doc: Document, sector_intel: Dict[ClientSector, SectorIntelligence], month: str
    ):
        """Create Page 2: Client Sector Intelligence"""

        # Header
        self._add_header(doc, "CLIENT SECTOR INTELLIGENCE", month)

        # Prioritize sectors with most intelligence
        sorted_sectors = sorted(sector_intel.items(), key=lambda x: len(x[1].items), reverse=True)

        # Add top sectors (limit to fit on one page)
        sectors_to_include = sorted_sectors[:4]

        for sector, intelligence in sectors_to_include:
            if intelligence.items:  # Only include if there's actual intelligence
                self._add_sector_section(doc, sector, intelligence)

    def _add_header(self, doc: Document, title: str, subtitle: str):
        """Add a professional header to the document with logo"""
        # Add logo if it exists
        if self.logo_path.exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(self.logo_path), width=Inches(2.0))
                p.paragraph_format.space_after = Pt(self.spacing["table_after"])
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"Could not add logo to document: {e}")

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = self.ergo_colors["primary_blue"]
        # Subtitle (Month Year)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = self.ergo_colors["secondary_blue"]
        # Add a subtle line using orange accent
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 50).font.color.rgb = self.ergo_colors["orange"]
    def _add_section_heading(self, doc: Document, heading: str):
        """Add a section heading"""
        p = doc.add_paragraph()
        p.style = "CustomHeading2"
        p.add_run(heading)

    def _add_executive_summary(self, doc: Document, items: List[IntelligenceItem]):
        """Generate sharp, actionable executive summary in Ergo analytical voice"""

        # Extract analytical judgments from top items
        insights = self._extract_analytical_insights(items)

        # Structure: BOTTOM LINE (2 max) + KEY FINDINGS (3-5) + WATCH FACTORS (2-3)
        bottom_line_items = insights.get("bottom_line", [])[:2]
        key_findings = insights.get("key_findings", [])[:5]
        watch_factors = insights.get("watch_factors", [])[:3]

        # Add BOTTOM LINE section if items exist
        if bottom_line_items:
            # Section header
            p = doc.add_paragraph()
            run = p.add_run("BOTTOM LINE")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.ergo_colors["primary_blue"]
            p.space_after = Pt(self.spacing["header_after"])
            for item_text in bottom_line_items:
                p = doc.add_paragraph()
                p.add_run(item_text)
                p.paragraph_format.space_after = Pt(self.spacing["paragraph"])
        # Add KEY FINDINGS section
        if key_findings:
            p = doc.add_paragraph()
            run = p.add_run("KEY FINDINGS")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.ergo_colors["primary_blue"]
            p.space_after = Pt(self.spacing["header_after"])
            for finding in key_findings:
                # Extract subheader and content from structured finding
                if isinstance(finding, dict):
                    subheader = finding.get("subheader", "")
                    content = finding.get("content", "")
                    bullets = finding.get("bullets", [])
                else:
                    # Legacy format - parse or use as-is
                    subheader, content, bullets = self._parse_key_finding(finding)

                # Add subheader (bold, slightly smaller than section header)
                if subheader:
                    p = doc.add_paragraph()
                    run = p.add_run(subheader)
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.ergo_colors["secondary_blue"]
                    p.paragraph_format.space_before = Pt(self.spacing["paragraph"])
                    p.paragraph_format.space_after = Pt(self.spacing["bullet"])
                # Add main paragraph content
                if content:
                    p = doc.add_paragraph()
                    p.add_run(content)
                    p.paragraph_format.space_after = Pt(self.spacing["bullet"])
                # Add supporting bullets
                for bullet in bullets:
                    p = doc.add_paragraph()
                    p.add_run(f"• {bullet}")
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(self.spacing["bullet"])

                # Add spacing after each finding
                if bullets:
                    doc.add_paragraph().paragraph_format.space_after = Pt(self.spacing["bullet"])
        # Add WATCH FACTORS section as a table
        if watch_factors:
            p = doc.add_paragraph()
            run = p.add_run("WATCH FACTORS")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.ergo_colors["primary_blue"]
            p.space_after = Pt(self.spacing["header_after"])
            # Create table with columns: Indicator | Current Status | What to Watch | Why It Matters
            # Ensure at least 5 rows for comprehensive watch factors
            num_rows = max(5, len(watch_factors))
            table = doc.add_table(rows=num_rows + 1, cols=4)  # +1 for header row
            table.style = "Table Grid"
            # Set column widths for 4 columns
            for row in table.rows:
                row.cells[0].width = Inches(1.5)
                row.cells[1].width = Inches(1.5)
                row.cells[2].width = Inches(2.0)
                row.cells[3].width = Inches(1.8)

            # Header row
            header_cells = table.rows[0].cells
            headers = ["Indicator", "Current Status", "What to Watch", "Why It Matters"]
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                # Style header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = self.ergo_colors["white"]
                # Set header background color
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B2946"/>')
                header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

            # Helper function to set cell text with consistent font styling
            def set_cell_text_with_style(cell, text, font_size=9):
                """Set cell text and apply consistent font styling"""
                # Clear existing paragraphs
                cell.text = ""
                # Add text as a run so we can style it
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(font_size)

            # Data rows
            for row_idx, factor in enumerate(watch_factors[:num_rows]):
                data_row = table.rows[row_idx + 1]  # +1 to skip header

                # Parse watch factor into components (now 4 fields)
                if isinstance(factor, dict):
                    indicator = factor.get("indicator", "")
                    current_status = factor.get("current_status", "Monitoring")
                    what_to_watch = factor.get("what_to_watch", "")
                    why_it_matters = factor.get("why_it_matters", "")
                else:
                    # Legacy format - parse the string
                    indicator, what_to_watch, why_it_matters = self._parse_watch_factor(factor)
                    current_status = "Monitoring"

                set_cell_text_with_style(data_row.cells[0], indicator)
                set_cell_text_with_style(data_row.cells[1], current_status)
                set_cell_text_with_style(data_row.cells[2], what_to_watch)
                set_cell_text_with_style(data_row.cells[3], why_it_matters)

            # Fill empty rows if we have fewer than 5 watch factors
            if len(watch_factors) < 5:
                for row_idx in range(len(watch_factors), 5):
                    data_row = table.rows[row_idx + 1]
                    # Generate placeholder content based on available intelligence
                    placeholder = self._generate_placeholder_watch_factor(row_idx, items)
                    if placeholder:
                        set_cell_text_with_style(data_row.cells[0], placeholder.get("indicator", "TBD"))
                        set_cell_text_with_style(data_row.cells[1], placeholder.get("current_status", "Monitoring"))
                        set_cell_text_with_style(data_row.cells[2], placeholder.get("what_to_watch", "Monitoring in progress"))
                        set_cell_text_with_style(data_row.cells[3], placeholder.get("why_it_matters", "Analysis pending"))
            # Add spacing after table
            doc.add_paragraph()

    def _extract_analytical_insights(self, items: List[IntelligenceItem]) -> dict:
        """Extract sharp analytical insights from intelligence items

        Note: Freshness filtering now handled globally in merge_intelligence_sources()
        to ensure consistency across entire report (not just Executive Summary)

        Uses AI-enhanced generation when available, with template fallback
        """
        # Try AI generation first if enabled
        if self.ai_generator:
            try:
                import logging

                logger = logging.getLogger(__name__)
                logger.info("Attempting AI-enhanced Executive Summary generation...")

                # Handle both running and new event loop scenarios
                try:
                    asyncio.get_running_loop()
                    # If there's a running loop, use a thread pool
                    import concurrent.futures
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        future = executor.submit(
                            asyncio.run,
                            self.ai_generator.generate_executive_summary(
                                items,
                                fallback_generator=lambda x: self._extract_insights_template(x),
                            ),
                        )
                        ai_summary = future.result(timeout=120)
                except RuntimeError:
                    # No running event loop - safe to use asyncio.run
                    ai_summary = asyncio.run(
                        self.ai_generator.generate_executive_summary(
                            items, fallback_generator=lambda x: self._extract_insights_template(x)
                        )
                    )

                if ai_summary and any(ai_summary.values()):
                    logger.info("✓ AI-enhanced Executive Summary generated successfully")
                    return ai_summary

            except Exception as e:
                import logging

                logging.getLogger(__name__).warning(f"AI generation failed, using template: {e}")

        # Fallback to template-based extraction
        return self._extract_insights_template(items)

    def _extract_insights_template(self, items: List[IntelligenceItem]) -> dict:
        """
        Template-based extraction (original implementation)
        Used as fallback when AI is unavailable or fails validation
        """
        bottom_line = []
        key_findings = []
        watch_factors = []

        # Sort by composite score (relevance × confidence)
        # No need to filter stale data here - already done in merge logic
        #
        # Prioritize ErgoMind for narrative leadership in Executive Summary
        # Ensures ErgoMind narrative items appear first in BOTTOM LINE and KEY FINDINGS
        ergomind_items = [i for i in items if i.source_type == "ergomind"]
        other_items = [i for i in items if i.source_type != "ergomind"]

        # Sort each group by composite score
        ergomind_sorted = sorted(
            ergomind_items, key=lambda x: (x.relevance_score * x.confidence), reverse=True
        )
        other_sorted = sorted(
            other_items, key=lambda x: (x.relevance_score * x.confidence), reverse=True
        )

        # Prioritize ErgoMind: scan ErgoMind first (~20 items), then data sources for supporting evidence
        prioritized = ergomind_sorted + other_sorted

        # Track content to avoid duplication
        seen_themes = set()

        for item in prioritized[:30]:  # Review top 30 items
            text = item.processed_content.lower()
            so_what = item.so_what_statement

            # Create theme fingerprint for deduplication
            theme = self._extract_theme(text, so_what)
            if theme in seen_themes:
                continue
            seen_themes.add(theme)

            # BOTTOM LINE: Immediate threats requiring action
            if (
                any(
                    kw in text
                    for kw in ["sanction", "blocking", "full blocking", "export control", "ban"]
                )
                and item.relevance_score > 0.7
                and len(bottom_line) < 2
            ):
                insight = self._craft_bottom_line_statement(item)
                if insight:
                    bottom_line.append(insight)
                    continue

            # KEY FINDINGS: Specific, actionable intelligence
            if item.relevance_score > 0.6 and len(key_findings) < 5:
                insight = self._craft_key_finding_statement(item)
                if insight and insight not in key_findings:
                    key_findings.append(insight)
                    continue

            # WATCH FACTORS: Forward-looking indicators
            if (
                any(
                    kw in text
                    for kw in ["trend", "signal", "increase", "decrease", "shift", "emerging"]
                )
                and len(watch_factors) < 3
            ):
                insight = self._craft_watch_factor_statement(item)
                if insight and insight not in watch_factors:
                    watch_factors.append(insight)

        return {
            "bottom_line": bottom_line,
            "key_findings": key_findings,
            "watch_factors": watch_factors,
        }

    def _extract_theme(self, text: str, so_what: str) -> str:
        """Extract core theme for deduplication"""
        # Combine key words from both text and so_what
        combined = (text + " " + so_what.lower())[:200]

        # Extract theme based on key terms
        if "sanction" in combined:
            return "sanctions"
        elif "tariff" in combined or "trade" in combined:
            return "trade"
        elif "inflation" in combined or "interest rate" in combined:
            return "monetary"
        elif "china" in combined and "us" in combined:
            return "us-china"
        elif "russia" in combined:
            return "russia"
        elif "aviation" in combined:
            return "aviation-specific"
        elif "fuel" in combined or "oil" in combined:
            return "fuel-costs"
        elif "technology" in combined or "semiconductor" in combined:
            return "tech-sector"
        else:
            # Use first 30 chars as fallback theme
            return combined[:30].strip()

    def _craft_bottom_line_statement(self, item: IntelligenceItem) -> str:
        """Craft authoritative bottom-line statement"""
        text = item.processed_content.lower()

        # Extract specific, high-impact statements
        if "rosneft" in text or "lukoil" in text or ("russian oil" in text and "sanction" in text):
            return "US full blocking sanctions on Russian oil majors create immediate compliance risk for clients with Russian exposure or routes through sanctioned jurisdictions."

        if "export control" in text and (
            "china" in text or "semiconductor" in text or "technology" in text
        ):
            return "Expanded US export controls on advanced technology require immediate client screening for Silicon Valley and tech sector exposure to restricted jurisdictions."

        if "full blocking" in text or "sdn list" in text:
            return "New sanctions regime requires client exposure audit within 30 days - particularly technology, energy, and financial sector clients with international operations."

        # Dynamic extraction for other high-priority threats
        if "airspace closure" in text or "flight ban" in text:
            return "New airspace restrictions require immediate route review and client notification - operational impacts within 48-72 hours."

        if "compliance risk" in text and "immediate" in text:
            # Extract the specific risk if mentioned
            sentences = item.processed_content.split(".")
            for sent in sentences[:3]:
                if "compliance" in sent.lower() and len(sent) > 40:
                    return sent.strip() + "."

        return None

    def _craft_key_finding_statement(self, item: IntelligenceItem) -> str:
        """Craft analytical key finding in Ergo voice - lead with judgment, support with evidence

        Always returns content - uses processed_content as fallback if no specific pattern matches
        """
        text = item.processed_content.lower()
        original_text = item.processed_content

        # Ergo analytical patterns: "Ergo assesses...", "Ergo believes...", probability statements

        # US-China relations
        if "truce" in text and ("october" in text or "china" in text):
            return "Ergo assesses that the US-China strategic truce (October 30) will hold through Q1 2026, stabilizing cross-Pacific routing. However, underlying export controls on semiconductors and advanced technology remain in force, creating compliance burdens for technology sector clients with Asia-Pacific operations."

        # Inflation/monetary policy
        if "inflation" in text and any(d.isdigit() for d in text):
            if "3.5" in text or "3.5%" in text:
                return "Ergo expects US inflation to remain above 3.5% through 2025, driving continued Federal Reserve rate increases. Aircraft financing costs will likely rise 15-25 basis points per quarter, affecting charter operators' cost structures and client pricing."

        # Tariff/trade litigation
        if "supreme court" in text and "tariff" in text:
            return "Ergo believes Supreme Court review of IEEPA-based tariffs creates significant regulatory uncertainty for Q1 2026. A ruling against executive tariff authority could trigger retroactive refunds and destabilize cost projections for clients in manufacturing and trade-sensitive sectors."

        # Russia sanctions
        if ("rosneft" in text or "lukoil" in text or "russian oil" in text) and "sanction" in text:
            return "Ergo assesses that US full blocking sanctions on Russian oil majors (Rosneft, Lukoil) create immediate compliance risk for energy sector clients. Operators must audit client exposure to sanctioned entities and routes through affected jurisdictions within 30 days."

        # China stimulus
        if "china" in text and ("stimulus" in text or "trillion" in text):
            # Extract specific amount if present
            amount_match = None
            if "$3" in text or "3 trillion" in text:
                amount_match = "$3+ trillion"

            if amount_match:
                return f"Ergo judges that China's {amount_match} stimulus package will stabilize domestic demand through mid-2026, supporting business aviation activity for clients with mainland operations. However, export-oriented sectors face persistent uncertainty from ongoing US-China tech restrictions."

        # EU fiscal constraints
        if "eu" in text and ("loan" in text or "fiscal" in text or "140" in text):
            return "Ergo assesses that EU fiscal constraints, evidenced by delays in the €140 billion Ukraine loan scheme, signal tightening corporate travel budgets across European clients in 2026. Public sector and cost-sensitive private clients will likely reduce discretionary aviation expenditures."

        # Fallback: Extract analytical sentences from original content
        sentences = original_text.split(".")
        for sent in sentences[:8]:
            sent_lower = sent.lower()
            # Look for Ergo assessments or expert-sourced statements
            if (
                "ergo assess" in sent_lower
                or "ergo believes" in sent_lower
                or "ergo expects" in sent_lower
                or "ergo judges" in sent_lower
            ):
                if len(sent) > 80 and len(sent) < 300:
                    return sent.strip() + "."

            # Look for probabilistic analytical statements
            if (
                any(
                    kw in sent_lower
                    for kw in ["likely", "unlikely", "probable", "expects", "believes"]
                )
                and len(sent) > 80
                and len(sent) < 250
            ):
                return sent.strip() + "."

        # ALWAYS return content - use the processed content directly if no specific pattern matches
        # Truncate at sentence boundary if needed
        if len(original_text) > 300:
            sentences = original_text.split(". ")
            truncated = ""
            for sent in sentences:
                if len(truncated) + len(sent) + 2 < 280:
                    truncated += sent + ". "
                else:
                    break
            if truncated:
                return truncated.strip()
            else:
                return sentences[0] + "." if sentences else original_text[:280]

        return original_text

    def _craft_watch_factor_statement(self, item: IntelligenceItem) -> str:
        """Craft forward-looking watch factor - always returns content"""
        text = item.processed_content.lower()
        original_text = item.processed_content

        if "data center" in text:
            return "Monitor corporate data center expansion in Middle America - signals sustained business aviation demand for project deployment and executive travel."

        if "eu" in text and "loan" in text:
            return "Track EU loan scheme execution (€140B target) - delays indicate fiscal constraints affecting European corporate travel budgets."

        if "supply chain" in text and ("china" in text or "asia" in text):
            return "Watch US-China supply chain decoupling velocity - accelerating separation drives increased site visit requirements for tech and manufacturing clients."

        if "inflation" in text or "interest rate" in text:
            return "Monitor Federal Reserve rate decisions and inflation trajectory - direct impact on aircraft financing costs and client budget decisions."

        if "sanction" in text:
            return "Track OFAC sanctions updates and enforcement actions - compliance requirements may change rapidly for international operations."

        if "china" in text:
            return "Watch US-China policy developments and diplomatic signals - sets tone for cross-Pacific business climate and travel demand."

        if "oil" in text or "fuel" in text or "energy" in text:
            return "Monitor crude oil and jet fuel price volatility - direct impact on operating costs and charter pricing."

        if "tariff" in text or "trade" in text:
            return "Track trade policy developments and tariff announcements - affects supply chain patterns and client travel requirements."

        if "technology" in text or "semiconductor" in text:
            return "Monitor tech sector export controls and regulatory changes - impacts Silicon Valley client compliance and travel patterns."

        # Fallback: Create generic watch factor from content
        # Extract first sentence and format as watch factor
        first_sentence = original_text.split(".")[0] if "." in original_text else original_text
        if len(first_sentence) > 80:
            first_sentence = first_sentence[:80]
        return (
            f"Monitor: {first_sentence} - potential implications for business aviation operations."
        )
    def _parse_key_finding(self, finding_text: str) -> tuple:
        """
        Parse a legacy key finding string into structured components.
        Returns (subheader, content, bullets) tuple.
        """
        text = finding_text.strip()

        # Try to extract a thematic subheader based on content analysis
        text_lower = text.lower()

        # Determine subheader based on key themes
        if "sanction" in text_lower or "blocking" in text_lower:
            subheader = "Sanctions & compliance risk"
        elif "china" in text_lower and (
            "us" in text_lower or "trade" in text_lower or "truce" in text_lower
        ):
            subheader = "US-China relations"
        elif (
            "inflation" in text_lower
            or "interest rate" in text_lower
            or "federal reserve" in text_lower
        ):
            subheader = "Monetary policy & financing"
        elif "tariff" in text_lower or "trade" in text_lower:
            subheader = "Trade policy & tariffs"
        elif (
            "technology" in text_lower
            or "semiconductor" in text_lower
            or "export control" in text_lower
        ):
            subheader = "Technology sector restrictions"
        elif "eu" in text_lower or "europe" in text_lower:
            subheader = "European regulatory landscape"
        elif "stimulus" in text_lower or "fiscal" in text_lower:
            subheader = "Fiscal policy & economic stimulus"
        elif "aviation" in text_lower or "aircraft" in text_lower or "fuel" in text_lower:
            subheader = "Aviation industry outlook"
        elif "regulation" in text_lower or "compliance" in text_lower or "policy" in text_lower:
            subheader = "Regulatory developments"
        else:
            # Extract first key phrase as subheader
            subheader = "Strategic development"

        # Split content into main paragraph and supporting bullets
        # Look for sentence boundaries to create bullets
        sentences = text.split(". ")

        if len(sentences) >= 3:
            # First 1-2 sentences as main content
            content = ". ".join(sentences[:2]) + "."
            # Remaining sentences as bullets
            bullets = [
                s.strip() + ("." if not s.strip().endswith(".") else "")
                for s in sentences[2:]
                if s.strip()
            ]
        elif len(sentences) == 2:
            content = sentences[0] + "."
            bullets = [
                sentences[1].strip() + ("." if not sentences[1].strip().endswith(".") else "")
            ]
        else:
            content = text
            bullets = []

        # Generate contextual bullets if we don't have enough
        if len(bullets) < 2:
            bullets.extend(
                self._generate_contextual_bullets(text_lower, subheader, 2 - len(bullets))
            )
        return (subheader, content, bullets[:3])  # Limit to 3 bullets max

    def _generate_contextual_bullets(self, text_lower: str, subheader: str, count: int) -> list:
        """Generate contextual supporting bullets based on content theme"""
        bullets = []

        if "sanction" in text_lower or "compliance" in subheader.lower():
            bullets.extend(
                [
                    "Immediate client exposure audit recommended for affected jurisdictions.",
                    "Operations teams should review routing through sanctioned airspace or territories.",
                    "Legal counsel consultation advised for high-exposure clients.",
                ]
            )
        elif "china" in text_lower or "us-china" in subheader.lower():
            bullets.extend(
                [
                    "Technology sector clients face heightened compliance requirements for cross-border operations.",
                    "Routing adjustments may be required for Asia-Pacific destinations.",
                    "Monitor for escalation signals that could affect bilateral flight agreements.",
                ]
            )
        elif "inflation" in text_lower or "monetary" in subheader.lower():
            bullets.extend(
                [
                    "Aircraft financing costs expected to increase, affecting lease and purchase decisions.",
                    "Operating cost increases may require pricing adjustments for charter services.",
                    "Budget-conscious clients may reduce discretionary travel frequency.",
                ]
            )
        elif "tariff" in text_lower or "trade" in subheader.lower():
            bullets.extend(
                [
                    "Manufacturing sector clients may adjust site visit patterns based on supply chain shifts.",
                    "Import-dependent businesses face cost pressures affecting travel budgets.",
                    "Monitor for retaliatory measures affecting additional sectors.",
                ]
            )
        elif "technology" in text_lower or subheader.lower() == "technology sector restrictions":
            bullets.extend(
                [
                    "Silicon Valley and tech corridor clients require enhanced compliance screening.",
                    "Data sovereignty regulations may affect international operations.",
                    "Export control awareness briefings recommended for affected clients.",
                ]
            )
        else:
            bullets.extend(
                [
                    "Continuous monitoring recommended for operational implications.",
                    "Client briefings should address sector-specific exposure.",
                    "Coordinate with operations team on contingency planning.",
                ]
            )
        return bullets[:count]

    def _parse_watch_factor(self, factor_text: str) -> tuple:
        """
        Parse a legacy watch factor string into table components.
        Returns (indicator, what_to_watch, why_it_matters) tuple.
        """
        text = factor_text.strip()
        text_lower = text.lower()

        # Extract indicator keyword
        if "data center" in text_lower:
            indicator = "Data Center Expansion"
            what_to_watch = "Corporate data center construction activity in Middle America"
            why_it_matters = "Signals sustained business aviation demand for project deployment and executive travel"
        elif "eu" in text_lower and "loan" in text_lower:
            indicator = "EU Fiscal Execution"
            what_to_watch = "EU loan scheme execution progress (€140B target)"
            why_it_matters = (
                "Delays indicate fiscal constraints affecting European corporate travel budgets"
            )
        elif "supply chain" in text_lower:
            indicator = "Supply Chain Decoupling"
            what_to_watch = "US-China supply chain separation velocity"
            why_it_matters = (
                "Accelerating separation drives increased site visit requirements for tech clients"
            )
        elif "inflation" in text_lower or "interest" in text_lower:
            indicator = "Interest Rate Trajectory"
            what_to_watch = "Federal Reserve rate decisions and inflation data"
            why_it_matters = "Affects aircraft financing costs and client budget decisions"
        elif "sanction" in text_lower:
            indicator = "Sanctions Enforcement"
            what_to_watch = "OFAC enforcement actions and SDN list updates"
            why_it_matters = (
                "Compliance requirements may change rapidly for international operations"
            )
        elif "china" in text_lower:
            indicator = "China Policy Signals"
            what_to_watch = "Beijing policy announcements and US diplomatic engagement"
            why_it_matters = "Sets tone for cross-Pacific business climate and travel demand"
        elif "oil" in text_lower or "fuel" in text_lower:
            indicator = "Fuel Price Volatility"
            what_to_watch = "Crude oil prices and refining capacity indicators"
            why_it_matters = "Direct impact on operating costs and charter pricing"
        else:
            # Generic parsing - try to split on common delimiters
            if " - " in text:
                parts = text.split(" - ", 1)
                indicator = parts[0][:30]  # Truncate indicator
                remainder = parts[1] if len(parts) > 1 else ""
                what_to_watch = remainder[:80] if remainder else "Monitor for developments"
                why_it_matters = "Potential impact on business aviation operations"
            else:
                indicator = text[:25] + "..." if len(text) > 25 else text
                what_to_watch = text[:80] if len(text) > 80 else text
                why_it_matters = "Relevance to aviation sector operations"

        return (indicator, what_to_watch, why_it_matters)

    def _generate_placeholder_watch_factor(self, index: int, items: List[IntelligenceItem]) -> dict:
        """Generate placeholder watch factors when fewer than 5 are available.

        Returns uniform placeholders covering 5 key categories:
        1. Fuel & Energy Costs
        2. Interest Rates & Financing
        3. Trade Policy & Tariffs
        4. Geopolitical Stability
        5. Regulatory Environment
        """
        placeholders = [
            {
                "indicator": "Fuel & Energy Costs",
                "current_status": "Elevated volatility",
                "what_to_watch": "Jet fuel spot prices (Gulf Coast), crude oil futures, refinery capacity utilization",
                "why_it_matters": "Direct operating cost impact; 15-25% of charter costs; affects pricing strategy",
            },
            {
                "indicator": "Interest Rates & Financing",
                "current_status": "Rates holding steady",
                "what_to_watch": "Fed funds rate trajectory, 10-year Treasury yields, aircraft financing spreads",
                "why_it_matters": "Fleet acquisition costs, client capital allocation decisions, M&A activity levels",
            },
            {
                "indicator": "Trade Policy & Tariffs",
                "current_status": "Active negotiations",
                "what_to_watch": "Section 301 tariffs, bilateral trade agreements, supply chain restrictions",
                "why_it_matters": "Client sector exposure varies; tech and manufacturing clients most sensitive",
            },
            {
                "indicator": "Geopolitical Stability",
                "current_status": "Monitoring hotspots",
                "what_to_watch": "Regional conflict indicators, diplomatic developments, travel advisories",
                "why_it_matters": "Route availability, client travel patterns, security planning requirements",
            },
            {
                "indicator": "Regulatory Environment",
                "current_status": "Evolving requirements",
                "what_to_watch": "FAA/EASA rule changes, environmental mandates, international flight permissions",
                "why_it_matters": "Compliance costs, operational flexibility, fleet planning implications",
            },
        ]

        return placeholders[index] if index < len(placeholders) else placeholders[0]

    def _add_economic_indicators_table(self, doc: Document, items: List[IntelligenceItem]):
        """Add economic indicators as a professional table"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        # Filter to get FRED economic data items
        fred_items = [item for item in items if item.source_type == "fred"]
        # If no FRED items, use any economic items
        if not fred_items:
            fred_items = items[:5]

        if not fred_items:
            p = doc.add_paragraph("No economic indicator data available for this period.")
            p.paragraph_format.space_after = Pt(self.spacing["header_after"])
            return

        # Create table with 4 columns: Indicator, Current Value, Trend, Aviation Impact
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        widths = [Inches(1.8), Inches(1.2), Inches(1.0), Inches(2.5)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Header row
        header_cells = table.rows[0].cells
        headers = ["Indicator", "Current Value", "Trend", "Aviation Impact"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Style header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            # Add shading to header
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
            header_cells[i]._tc.get_or_add_tcPr().append(shading)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Add data rows
        for item in fred_items[:5]:  # Limit to 5 indicators
            row = table.add_row()
            cells = row.cells

            # Extract indicator name from content
            indicator_name = self._extract_indicator_name(item)
            cells[0].text = indicator_name

            # Extract current value
            current_value = self._extract_value(item)
            cells[1].text = current_value

            # Determine trend
            trend = self._determine_trend(item)
            cells[2].text = trend

            # Aviation impact - use so_what statement or generate one
            impact = (
                item.so_what_statement
                if item.so_what_statement
                else self._generate_economic_impact(item)
            )
            # Truncate if too long
            if len(impact) > 120:
                impact = impact[:117] + "..."
            cells[3].text = impact

            # Style data cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Add spacing after table
        doc.add_paragraph()

    def _strip_markdown(self, text: str) -> str:
        """Remove markdown formatting artifacts from text for clean document rendering"""
        import re

        if not text:
            return text

        # Remove bold/italic markers (* and _)
        text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)  # ***bold*** -> bold
        text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)  # ___italic___ -> italic

        # Remove markdown headers (# ## ### etc)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

        # Remove markdown links [text](url) -> text
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

        # Remove markdown images ![alt](url)
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)

        # Remove bullet point markers at start of lines
        text = re.sub(r"^[\s]*[-*+]\s+", "", text, flags=re.MULTILINE)

        # Remove numbered list markers
        text = re.sub(r"^[\s]*\d+\.\s+", "", text, flags=re.MULTILINE)

        # Remove code backticks
        text = re.sub(r"`{1,3}([^`]+)`{1,3}", r"\1", text)

        # Remove blockquote markers
        text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)

        # Replace en-dash and em-dash with regular hyphen or remove
        text = text.replace("–", "-").replace("—", "-")

        # Remove bracketed numbers like [1], [2], (1), (2) - citation artifacts
        text = re.sub(r"\[\d+\]", "", text)
        text = re.sub(r"\(\d+\)", "", text)

        # Remove standalone numbers in brackets at word boundaries
        text = re.sub(r"\s*\[\d+(?:,\s*\d+)*\]\s*", " ", text)

        # Remove parenthetical asides that are just numbers or short codes
        text = re.sub(r"\s*\([A-Z]{2,4}\d*\)\s*", " ", text)

        # Remove HTML tags if any
        text = re.sub(r"<[^>]+>", "", text)

        # Fix common grammar issues
        # Remove orphaned punctuation
        text = re.sub(r"\s+([,;:])", r"\1", text)
        text = re.sub(r"([,;:])\s*([,;:])", r"\1", text)

        # Fix double periods
        text = re.sub(r"\.\.+", ".", text)

        # Fix spacing around hyphens used as dashes (keep compound words intact)
        text = re.sub(r"\s+-\s+", " - ", text)

        # Clean up extra whitespace
        text = re.sub(r"\s+", " ", text)
        text = text.strip()

        return text

    def _extract_indicator_name(self, item: IntelligenceItem) -> str:
        """Extract a clean indicator name from the item"""
        content = item.processed_content.lower()

        # Try to identify common economic indicators
        if "gdp" in content:
            return "GDP Growth"
        elif "inflation" in content or "cpi" in content:
            return "Inflation Rate"
        elif "unemployment" in content or "jobs" in content:
            return "Unemployment"
        elif "interest rate" in content or "fed" in content or "federal reserve" in content:
            return "Interest Rate"
        elif "oil" in content or "petroleum" in content or "energy" in content:
            return "Oil/Energy"
        elif "jet fuel" in content or "fuel" in content:
            return "Jet Fuel"
        elif "dollar" in content or "currency" in content or "exchange" in content:
            return "USD Index"
        elif "consumer" in content or "spending" in content:
            return "Consumer Spending"
        elif "manufacturing" in content or "pmi" in content:
            return "Manufacturing PMI"
        elif "trade" in content or "export" in content or "import" in content:
            return "Trade Balance"
        else:
            # Use first few words of category or content
            words = item.processed_content.split()[:3]
            return " ".join(words)[:25]
    def _extract_value(self, item: IntelligenceItem) -> str:
        """Extract the current value from content"""
        import re

        content = item.processed_content

        # Look for percentage patterns
        pct_match = re.search(r"(\d+\.?\d*)\s*%", content)
        if pct_match:
            return f"{pct_match.group(1)}%"

        # Look for dollar amounts
        dollar_match = re.search(
            r"\$(\d+\.?\d*)\s*(billion|million|trillion)?", content, re.IGNORECASE
        )
        if dollar_match:
            amount = dollar_match.group(1)
            unit = dollar_match.group(2) or ""
            return f"${amount}{unit[0].upper() if unit else ''}"

        # Look for just numbers with context
        num_match = re.search(r"(\d+\.?\d*)\s*(points?|bps|basis points)?", content)
        if num_match:
            return num_match.group(0)

        return "See details"

    def _determine_trend(self, item: IntelligenceItem) -> str:
        """Determine trend direction from content"""
        content = item.processed_content.lower()

        up_words = ["increase", "rose", "rising", "grew", "growth", "higher", "up", "gain", "surge"]
        down_words = ["decrease", "fell", "falling", "decline", "lower", "down", "drop", "slump"]
        stable_words = ["stable", "unchanged", "flat", "steady", "maintained"]
        if any(word in content for word in up_words):
            return "Rising"
        elif any(word in content for word in down_words):
            return "Falling"
        elif any(word in content for word in stable_words):
            return "Stable"
        else:
            return "Monitor"

    def _generate_economic_impact(self, item: IntelligenceItem) -> str:
        """Generate an aviation impact statement for economic data"""
        content = item.processed_content.lower()

        if "fuel" in content or "oil" in content or "energy" in content:
            return "Directly affects operating costs and charter pricing"
        elif "interest" in content or "fed" in content:
            return "Influences aircraft financing and client capital allocation"
        elif "inflation" in content:
            return "Impacts operating margins and pricing strategy"
        elif "gdp" in content or "growth" in content:
            return "Indicator of business travel demand trajectory"
        elif "unemployment" in content or "jobs" in content:
            return "Correlates with corporate travel budgets"
        elif "consumer" in content or "spending" in content:
            return "Signals leisure charter demand patterns"
        elif "dollar" in content or "currency" in content:
            return "Affects international charter competitiveness"
        elif "trade" in content:
            return "Influences cross-border business travel needs"
        else:
            return "Monitor for potential aviation sector impact"

    def _add_intelligence_items(self, doc: Document, items: List[IntelligenceItem]):
        """Add intelligence items with smart truncation at sentence boundaries"""
        for item in items:
            if item.relevance_score < 0.5:  # Skip low relevance items
                continue

            # Main finding - clean paragraph without bullets
            p = doc.add_paragraph()

            # Smart truncation at sentence boundaries - strip markdown artifacts
            content = self._strip_markdown(item.processed_content)
            if len(content) > 400:
                # Try to truncate at sentence boundary
                sentences = content.split(". ")
                truncated = ""
                for sentence in sentences:
                    # Check if adding this sentence would exceed limit
                    if len(truncated) + len(sentence) + 2 < 380:  # Leave room for ". "
                        truncated += sentence + ". "
                    else:
                        break

                # Only use complete sentences - no ellipsis
                if truncated and len(truncated) > 100:
                    content = truncated.rstrip()
                else:
                    # Fallback: extract first sentence only
                    first_period = content.find(". ")
                    if first_period > 50:  # Reasonable sentence length
                        content = content[: first_period + 1]
                    else:
                        # If no good sentence boundary, take first 300 chars at word boundary
                        content = content[:300]
                        last_space = content.rfind(" ")
                        if last_space > 250:
                            content = content[:last_space] + "."

            p.add_run(content)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])
            # Add source citation - differentiate between ErgoMind and GTA
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])
            if item.source_type == "gta":
                # GTA source with intervention ID and dates
                source_text = "   📊 GTA Data"
                if item.gta_intervention_id:
                    source_text += f" (ID: {item.gta_intervention_id})"
                if item.date_implemented:
                    source_text += f" | Implemented: {item.date_implemented}"

                run = p.add_run(source_text)
                run.font.size = Pt(9)
                run.font.color.rgb = self.ergo_colors["secondary_blue"]  # Blue for GTA
                run.font.italic = True

                # Add geographic data if available
                if item.gta_implementing_countries or item.gta_affected_countries:
                    p = doc.add_paragraph()
                    geo_text = "   "
                    if item.gta_implementing_countries:
                        geo_text += (
                            f"Implementing: {', '.join(item.gta_implementing_countries[:3])}"
                        )
                    if item.gta_affected_countries:
                        if item.gta_implementing_countries:
                            geo_text += " | "
                        geo_text += f"Affected: {', '.join(item.gta_affected_countries[:3])}"

                    run = p.add_run(geo_text)
                    run.font.size = Pt(8)
                    run.font.color.rgb = self.ergo_colors["dark_gray"]
                    run.font.italic = True
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(self.spacing["bullet"])
            elif item.source_type == "fred":
                # FRED source with series ID and observation date
                source_text = "   📊 FRED Economic Data"
                if item.fred_series_id:
                    source_text += f" ({item.fred_series_id})"
                if item.fred_observation_date:
                    source_text += f" | {item.fred_observation_date}"

                run = p.add_run(source_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 100, 0)  # Green for FRED (Federal Reserve)
                run.font.italic = True
            else:
                # ErgoMind source (existing format)
                if item.sources:
                    source_count = len(item.sources)
                    source_text = f"   Sources: Based on {source_count} source{'s' if source_count > 1 else ''}"
                    run = p.add_run(source_text)
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.ergo_colors["dark_gray"]
                    run.font.italic = True

            # Add "So What" in italics - clean format without arrows
            if item.so_what_statement:
                p = doc.add_paragraph()
                run1 = p.add_run("Impact: ")
                run1.font.italic = True
                run1.font.size = Pt(10)
                run2 = p.add_run(item.so_what_statement)
                run2.font.italic = True
                run2.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(self.spacing["paragraph"])
    def _add_regional_assessment(self, doc: Document, regional_items: List[IntelligenceItem]):
        """Add regional risk assessments with rigorous analytical content matching document tone"""
        regions = {"North America": [], "Europe": [], "Asia-Pacific": [], "Middle East": []}

        # Track which items we've already assigned to prevent duplication
        assigned_items = set()

        # Categorize items by PRIMARY region only - check both category and content
        for item in regional_items:
            item_id = id(item)
            if item_id in assigned_items:
                continue

            text_lower = item.processed_content.lower()
            category_lower = item.category.lower()
            primary_region = None

            if item.source_type == "gta":
                countries = item.gta_implementing_countries + item.gta_affected_countries
                countries_text = " ".join(countries).lower()

                if any(c in countries_text for c in ["united states", "canada", "mexico"]):
                    primary_region = "North America"
                elif any(c in countries_text for c in ["china", "japan", "india", "korea", "singapore", "australia", "taiwan"]):
                    primary_region = "Asia-Pacific"
                elif any(c in countries_text for c in ["germany", "france", "uk", "united kingdom", "italy", "spain", "european union", "netherlands", "poland"]):
                    primary_region = "Europe"
                elif any(c in countries_text for c in ["saudi", "uae", "qatar", "israel", "iran", "kuwait", "bahrain"]):
                    primary_region = "Middle East"

            if not primary_region:
                if any(kw in text_lower or kw in category_lower for kw in ["united states", "u.s.", "america", "canada", "mexico", "washington", "federal"]):
                    primary_region = "North America"
                elif any(kw in text_lower or kw in category_lower for kw in ["china", "japan", "india", "asia", "pacific", "beijing", "tokyo", "seoul"]):
                    primary_region = "Asia-Pacific"
                elif any(kw in text_lower or kw in category_lower for kw in ["europe", "eu", "britain", "france", "germany", "uk", "brussels", "ecb"]):
                    primary_region = "Europe"
                elif any(kw in text_lower or kw in category_lower for kw in ["middle east", "saudi", "uae", "qatar", "israel", "gulf", "opec"]):
                    primary_region = "Middle East"

            if primary_region:
                regions[primary_region].append(item)
                assigned_items.add(item_id)

        # Generate comprehensive regional assessments
        for region, items in regions.items():
            if items:
                # Region header
                p = doc.add_paragraph()
                run = p.add_run(region)
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = self.ergo_colors["primary_blue"]
                p.paragraph_format.space_before = Pt(self.spacing["subsection_before"])
                p.paragraph_format.space_after = Pt(self.spacing["bullet"])

                # Sort by relevance and use top items
                sorted_items = sorted(items, key=lambda x: x.relevance_score, reverse=True)

                # Generate multi-paragraph assessment using top 2-3 items
                self._write_regional_assessment_paragraphs(doc, region, sorted_items[:3])

    def _write_regional_assessment_paragraphs(
        self, doc: Document, region: str, items: List[IntelligenceItem]
    ):
        """
        Write multi-paragraph regional assessments with rigorous analytical content.
        Matches the tone, depth, and structure of KEY FINDINGS sections.
        """
        import re

        if not items:
            return

        # Collect and clean all content from items
        all_content = []
        for item in items:
            content = self._strip_markdown(item.processed_content)
            if content:
                all_content.append(content)

        if not all_content:
            return

        # Extract specifics from combined content for analytical framing
        combined_text = " ".join(all_content)
        specifics = self._extract_specifics(combined_text, "")

        # Build the assessment as cohesive paragraphs
        paragraphs = []

        # PARAGRAPH 1: Lead with the most significant development
        primary_item = items[0]
        primary_content = self._strip_markdown(primary_item.processed_content)

        # Extract substantive sentences (filter out metadata and short fragments)
        sentences = [s.strip() for s in primary_content.split(". ") if s.strip()]
        substantive_sentences = []
        for sent in sentences:
            # Skip metadata prefixes
            if any(sent.startswith(prefix) for prefix in [
                "Baseline:", "Background:", "Context:", "Note:", "Source:",
                "Reference:", "Update:", "Summary:", "Overview:", "Key Point:",
                "Action:", "Status:", "**", "- ", "• "
            ]):
                continue
            # Skip very short or fragment sentences
            if len(sent) < 40 or not sent[0].isupper():
                continue
            substantive_sentences.append(sent)

        if substantive_sentences:
            # Use first 2-3 substantive sentences as lead paragraph
            lead_text = ". ".join(substantive_sentences[:3])
            if not lead_text.endswith("."):
                lead_text += "."
            paragraphs.append(lead_text)

        # PARAGRAPH 2: Supporting intelligence from secondary items
        if len(items) > 1:
            supporting_sentences = []
            for item in items[1:3]:
                content = self._strip_markdown(item.processed_content)
                item_sentences = [s.strip() for s in content.split(". ") if s.strip()]
                for sent in item_sentences:
                    if any(sent.startswith(prefix) for prefix in [
                        "Baseline:", "Background:", "Context:", "Note:", "Source:",
                        "Reference:", "Update:", "Summary:", "Overview:", "Key Point:",
                        "Action:", "Status:", "**", "- ", "• "
                    ]):
                        continue
                    if len(sent) >= 40 and sent[0].isupper():
                        # Check for data points (numbers, percentages, currencies)
                        has_data = bool(re.search(
                            r"\d+%|\$\d+|€\d+|£\d+|\d+\s*(?:billion|million|trillion)",
                            sent, re.IGNORECASE
                        ))
                        if has_data or len(supporting_sentences) < 2:
                            supporting_sentences.append(sent)
                            if len(supporting_sentences) >= 3:
                                break
                if len(supporting_sentences) >= 3:
                    break

            if supporting_sentences:
                support_text = ". ".join(supporting_sentences[:3])
                if not support_text.endswith("."):
                    support_text += "."
                paragraphs.append(support_text)

        # PARAGRAPH 3: Analytical implications (if we have enough material)
        if len(paragraphs) >= 1:
            # Build implications based on extracted data
            implications = []

            # Add quantitative context if available
            if specifics.get("percentages"):
                pct = specifics["percentages"][0]
                implications.append(f"measures involving {pct} adjustments")

            if specifics.get("currencies"):
                curr = specifics["currencies"][0]
                implications.append(f"financial exposure of {curr}")

            # Build operational implications sentence based on region and content themes
            content_lower = combined_text.lower()

            if "tariff" in content_lower or "trade" in content_lower:
                if region == "North America":
                    implications.append("supply chain routing decisions for cross-border operations")
                elif region == "Asia-Pacific":
                    implications.append("manufacturing sourcing and procurement strategies")
                elif region == "Europe":
                    implications.append("EU customs and compliance considerations")
                elif region == "Middle East":
                    implications.append("regional trade corridor planning")

            if "sanction" in content_lower or "restriction" in content_lower:
                implications.append("enhanced due diligence on counterparty relationships")

            if "fuel" in content_lower or "energy" in content_lower or "oil" in content_lower:
                implications.append("fuel cost hedging and route optimization")

            if "regulation" in content_lower or "regulatory" in content_lower:
                implications.append("compliance monitoring and reporting requirements")

            # Only add implications paragraph if we have substantive ones
            if implications and len(implications) >= 2:
                impl_text = f"For Solairus operations in {region}, these developments warrant attention to {', '.join(implications[:3])}."
                paragraphs.append(impl_text)

        # Write paragraphs to document with proper formatting
        for i, para_text in enumerate(paragraphs):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.name = self.body_font
            run.font.size = Pt(self.font_sizes["body"])

            # First paragraph: normal indent; subsequent: slight indent for flow
            if i > 0:
                p.paragraph_format.first_line_indent = Inches(0.25)

            p.paragraph_format.space_after = Pt(self.spacing["paragraph"])
            p.paragraph_format.line_spacing = 1.15

    def _craft_regional_assessment(
        self, region: str, primary_item: IntelligenceItem, all_items: List[IntelligenceItem]
    ) -> str:
        """Craft a specific, data-driven regional assessment using actual intelligence details"""
        import re

        # Strip markdown and clean the content
        content = self._strip_markdown(primary_item.processed_content)
        # Extract first sentence as a title-like summary (IntelligenceItem has no title field)
        first_sentence = ""
        if content:
            sentences = content.split(". ")
            if sentences:
                first_sentence = sentences[0].strip()

        # Extract specific details from the content
        specifics = self._extract_specifics(content, first_sentence)

        # Build a specific assessment based on extracted details
        assessment = self._build_specific_assessment(region, specifics, content, primary_item)

        return assessment

    def _extract_specifics(self, content: str, title: str) -> dict:
        """Extract specific data points from intelligence content"""
        import re

        combined = f"{title} {content}"
        specifics = {
            "percentages": [],
            "currencies": [],
            "countries": [],
            "companies": [],
            "dates": [],
            "policies": [],
            "numbers": [],
        }

        # Extract percentages (e.g., "25%", "10-15%", "up to 50%")
        pct_matches = re.findall(r"(\d+(?:\.\d+)?(?:-\d+(?:\.\d+)?)?%)", combined)
        specifics["percentages"] = pct_matches[:3]

        # Extract currency amounts (e.g., "$50 billion", "€100 million", "¥500 billion")
        currency_matches = re.findall(
            r"[\$€£¥]\s*\d+(?:\.\d+)?\s*(?:billion|million|trillion|bn|mn|B|M|T)?",
            combined,
            re.IGNORECASE,
        )
        specifics["currencies"] = currency_matches[:3]

        # Extract specific country names
        country_patterns = [
            "United States",
            "China",
            "Russia",
            "Japan",
            "Germany",
            "France",
            "United Kingdom",
            "UK",
            "India",
            "Brazil",
            "Canada",
            "Australia",
            "South Korea",
            "Mexico",
            "Italy",
            "Spain",
            "Saudi Arabia",
            "UAE",
            "Israel",
            "Iran",
            "Taiwan",
            "Singapore",
            "Hong Kong",
            "Indonesia",
            "Turkey",
            "Poland",
            "Netherlands",
            "Switzerland",
            "Sweden",
            "Norway",
        ]
        for country in country_patterns:
            if country.lower() in combined.lower():
                if country not in specifics["countries"]:
                    specifics["countries"].append(country)

        # Extract company names (capitalized multi-word phrases that look like companies)
        company_matches = re.findall(
            r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:\s+(?:Inc|Corp|Ltd|LLC|Co)\.?)?)\b", combined
        )
        for match in company_matches[:5]:
            if len(match) > 3 and match not in country_patterns:
                specifics["companies"].append(match)

        # Extract dates and timeframes
        date_matches = re.findall(
            r"(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}|Q[1-4]\s+\d{4}|(?:first|second|third|fourth)\s+quarter",
            combined,
            re.IGNORECASE,
        )
        specifics["dates"] = date_matches[:2]

        # Extract policy/regulation names (phrases in quotes or with specific keywords)
        policy_matches = re.findall(r'"([^"]+)"', combined)
        specifics["policies"] = [p for p in policy_matches if len(p) > 5 and len(p) < 80][:2]

        # Extract large numbers (millions, billions)
        num_matches = re.findall(
            r"\b(\d+(?:\.\d+)?\s*(?:million|billion|trillion))\b", combined, re.IGNORECASE
        )
        specifics["numbers"] = num_matches[:3]

        return specifics

    def _build_specific_assessment(
        self, region: str, specifics: dict, content: str, item: IntelligenceItem
    ) -> str:
        """Build a specific assessment using extracted data points - NEVER use generic templates"""
        import re

        content_lower = content.lower()

        # Metadata prefixes to filter out from sentences
        metadata_prefixes = [
            "Baseline:",
            "Background:",
            "Context:",
            "Note:",
            "Source:",
            "Reference:",
            "Update:",
            "Summary:",
            "Overview:",
            "Key Point:",
            "Action:",
            "Status:",
        ]

        def is_metadata_sentence(sentence: str) -> bool:
            """Check if sentence starts with a metadata prefix"""
            for prefix in metadata_prefixes:
                if sentence.startswith(prefix) or sentence.startswith(prefix.lower()):
                    return True
            return False

        # Extract all sentences for later use, filtering out metadata
        all_sentences = [
            s.strip() for s in content.split(". ")
            if s.strip() and not is_metadata_sentence(s.strip())
        ]

        # Get extracted specifics
        countries = specifics.get("countries", [])
        percentages = specifics.get("percentages", [])
        currencies = specifics.get("currencies", [])

        # PRIORITY 1: Find a sentence from content that contains actual data
        for sent in all_sentences:
            if len(sent) < 30 or len(sent) > 350:
                continue
            if sent and sent[0].islower():
                continue
            # Skip metadata-like sentences
            if is_metadata_sentence(sent):
                continue
            # Look for sentences with numbers, percentages, dollar amounts, or specific countries
            has_data = bool(
                re.search(r"\d+%|\$\d+|€\d+|£\d+|\d+\s*(?:billion|million|trillion)", sent, re.IGNORECASE)
            )
            has_country = any(c.lower() in sent.lower() for c in countries) if countries else False
            if has_data or has_country:
                if not sent.endswith("."):
                    sent = sent + "."
                return sent

        # PRIORITY 2: Build assessment from extracted specifics
        parts = []

        # Determine the primary topic and build specific sentence
        if "tariff" in content_lower:
            if percentages:
                tariff_rate = percentages[0]
                if countries:
                    parts.append(f"{tariff_rate} tariffs on {countries[0]} goods")
                else:
                    parts.append(f"New {tariff_rate} tariffs")
            elif countries:
                parts.append(f"Tariff measures targeting {countries[0]}")

        elif "sanction" in content_lower or "restriction" in content_lower:
            if countries:
                parts.append(f"Sanctions affecting {', '.join(countries[:2])}")
            if specifics.get("policies"):
                parts.append(f"under {specifics['policies'][0]}")

        elif "export control" in content_lower or "export restriction" in content_lower:
            if countries:
                parts.append(f"Export controls targeting {countries[0]}")
            if "semiconductor" in content_lower or "chip" in content_lower:
                parts.append("in the semiconductor sector")
            elif "technology" in content_lower:
                parts.append("on technology products")

        elif currencies:
            if "stimulus" in content_lower or "spending" in content_lower:
                parts.append(f"{currencies[0]} stimulus package")
            elif "investment" in content_lower:
                parts.append(f"{currencies[0]} investment")
            elif "deficit" in content_lower or "debt" in content_lower:
                parts.append(f"{currencies[0]} fiscal impact")

        elif percentages and ("growth" in content_lower or "gdp" in content_lower):
            parts.append(f"{percentages[0]} GDP growth")
            if countries:
                parts.append(f"in {countries[0]}")

        elif percentages and "inflation" in content_lower:
            parts.append(f"Inflation at {percentages[0]}")
            if countries:
                parts.append(f"in {countries[0]}")

        # If we have specific parts, build the assessment
        if parts:
            base = " ".join(parts)
            # Add aviation/business context
            if "tariff" in content_lower or "trade" in content_lower:
                assessment = (
                    f"{base} may affect supply chains and client travel to affected regions."
                )
            elif "sanction" in content_lower:
                assessment = f"{base} require compliance review for operations involving these jurisdictions."
            elif "growth" in content_lower or "stimulus" in content_lower:
                assessment = f"{base} signals potential increase in business aviation demand."
            elif "inflation" in content_lower or "recession" in content_lower:
                assessment = f"{base} may impact corporate travel budgets in the near term."
            else:
                assessment = f"{base} warrants monitoring for aviation operational impacts."

            # Ensure proper formatting
            if assessment[0].islower():
                assessment = assessment[0].upper() + assessment[1:]
            if not assessment.endswith("."):
                assessment = assessment + "."
            return assessment

        # PRIORITY 3: Try to extract ANY meaningful sentence from content (not just data-heavy ones)
        for sent in all_sentences:
            sent = sent.strip()
            if len(sent) < 40 or len(sent) > 300:
                continue
            if sent and sent[0].islower():
                continue
            # Accept any sentence that looks like a complete thought
            if not sent.endswith("."):
                sent = sent + "."
            return sent

        # PRIORITY 4: Use first sentence if it's minimally informative
        if all_sentences and len(all_sentences[0]) > 20:
            first = all_sentences[0].strip()
            if first and first[0].isupper():
                if not first.endswith("."):
                    first = first + "."
                return first

        # PRIORITY 5: Build from whatever specifics we have, even without topic context
        if countries:
            return f"Developments in {', '.join(countries[:2])} warrant monitoring for operational impacts."
        if percentages:
            return f"Economic indicators showing {percentages[0]} changes may affect regional operations."
        if currencies:
            return f"{currencies[0]} fiscal developments may impact business activity in the region."

        # Final fallback - generic but region-specific (avoid so_what_statement which has template text)
        fallbacks = {
            "North America": "Monitor US regulatory and trade policy developments for operational impact.",
            "Europe": "Track EU regulatory changes affecting aviation operations.",
            "Asia-Pacific": "Assess regional trade dynamics for impact on cross-Pacific operations.",
            "Middle East": "Monitor security and energy developments affecting regional operations.",
            "Latin America": "Track economic and political developments in key markets.",
            "Africa": "Monitor regulatory environment in emerging aviation markets.",
        }
        return fallbacks.get(region, "Continue monitoring regional developments.")
    def _add_regulatory_outlook(self, doc: Document, reg_items: List[IntelligenceItem]):
        """Add regulatory outlook section with actual content from items"""
        if not reg_items:
            p = doc.add_paragraph()
            p.add_run("No significant regulatory changes identified this period.")
            return

        # Sort by relevance and take top items
        sorted_items = sorted(reg_items, key=lambda x: x.relevance_score, reverse=True)

        # Add regulatory findings as bullet points
        added_content = set()
        items_added = 0

        for item in sorted_items[:6]:  # Check up to 6 items
            if items_added >= 4:  # Show max 4 items
                break

            content = self._strip_markdown(item.processed_content)

            # Skip empty or too short content
            if not content or len(content) < 50:
                continue

            # Skip duplicates
            content_key = content[:80].lower()
            if content_key in added_content:
                continue
            added_content.add(content_key)

            # Truncate to reasonable length at sentence boundary
            if len(content) > 200:
                sentences = content.split(". ")
                truncated = ""
                for sentence in sentences:
                    if len(truncated) + len(sentence) + 2 <= 180:
                        truncated += sentence + ". "
                    else:
                        break
                content = truncated.rstrip() if truncated else sentences[0][:180] + "..."

            p = doc.add_paragraph()
            p.add_run(f"• {content}")
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])
            items_added += 1

        # If no items were added from content, show fallback message
        if items_added == 0:
            p = doc.add_paragraph()
            p.add_run("No significant regulatory changes identified this period.")
            return

    def _add_sector_section(
        self, doc: Document, sector: ClientSector, intelligence: SectorIntelligence
    ):
        """Add a section for a specific client sector"""
        # Sector heading
        sector_name = sector.value.replace("_", " ").upper()
        if sector == ClientSector.TECHNOLOGY:
            sector_name = "TECHNOLOGY SECTOR CLIENTS"
        elif sector == ClientSector.FINANCE:
            sector_name = "FINANCIAL SECTOR CLIENTS"
        elif sector == ClientSector.REAL_ESTATE:
            sector_name = "REAL ESTATE SECTOR CLIENTS"
        elif sector == ClientSector.ENTERTAINMENT:
            sector_name = "ENTERTAINMENT & MEDIA SECTOR CLIENTS"

        self._add_section_heading(doc, sector_name)

        # Critical Developments
        p = doc.add_paragraph()
        run = p.add_run("Critical Developments:")
        run.font.bold = True
        run.font.color.rgb = self.ergo_colors["dark_gray"]

        # Add top 2 items for this sector - complete sentences only
        # Track what we've added to avoid duplicates
        added_content = set()
        items_added = 0

        for item in intelligence.items[:5]:  # Check up to 5 items to find 2 unique ones
            if items_added >= 2:
                break

            content = self._strip_markdown(item.processed_content)

            # Skip empty or too short content
            if not content or len(content) < 50:
                continue

            # Skip generic/filler content
            generic_phrases = [
                "you can already see which indicators",
                "what to watch for as",
                "will be most informative",
                "based on Ergo's current macro work",
            ]
            if any(phrase.lower() in content.lower() for phrase in generic_phrases):
                continue

            # Skip if we've already added very similar content
            content_key = content[:100].lower()
            if content_key in added_content:
                continue
            added_content.add(content_key)

            # Truncate to complete sentences if needed
            if len(content) > 250:
                sentences = content.split(". ")
                truncated = ""
                for sentence in sentences:
                    if len(truncated) + len(sentence) + 2 <= 250:
                        truncated += sentence + ". "
                    else:
                        break
                content = truncated.rstrip() if truncated else sentences[0] + "."

            p = doc.add_paragraph()
            p.add_run(f"• {content}")
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])
            items_added += 1
        # If no items were added, add a placeholder message
        if items_added == 0:
            p = doc.add_paragraph()
            p.add_run("No significant sector-specific developments this period.")
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])

        # Collect unique action items - filter out generic ones
        generic_patterns = [
            "Monitor situation and prepare briefing",
            "Update market intelligence briefings",
            "Conduct immediate review of affected routes and file alternative flight plans"
        ]

        all_actions = []
        for item in intelligence.items[:3]:  # Top 3 items only
            for action in item.action_items:
                # Skip generic actions
                is_generic = any(pattern.lower() in action.lower() for pattern in generic_patterns)
                if not is_generic and action not in all_actions:
                    all_actions.append(action)

        # If we filtered out everything, generate sector-specific actions
        if not all_actions:
            all_actions = self._generate_sector_specific_actions(sector, intelligence.items[:2])

        for action in all_actions[:3]:
            p = doc.add_paragraph()
            p.add_run(action)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(self.spacing["bullet"])
        # Add spacing between sectors
        doc.add_paragraph()

    def _generate_sector_specific_actions(
        self, sector: ClientSector, items: List[IntelligenceItem]
    ) -> List[str]:
        """Generate sector-specific action items when generic ones are filtered out"""
        actions = []

        if sector == ClientSector.TECHNOLOGY:
            actions = [
                "Proactively reach out to technology sector clients about export control impacts",
                "Monitor Silicon Valley travel patterns for early demand signals",
                "Brief tech executives on cross-border data sovereignty requirements",
            ]
        elif sector == ClientSector.FINANCE:
            actions = [
                "Schedule check-ins with financial sector clients to discuss market volatility impacts",
                "Position for relationship-critical travel during M&A activity periods",
                "Monitor capital flow restrictions affecting client budgets",
            ]
        elif sector == ClientSector.REAL_ESTATE:
            actions = [
                "Alert real estate clients to regional market developments affecting property tours",
                "Track construction material costs impact on development timelines",
                "Coordinate with clients on site visit scheduling around market conditions",
            ]
        elif sector == ClientSector.ENTERTAINMENT:
            actions = [
                "Coordinate with entertainment clients on production schedule changes",
                "Monitor content regulation impacts on international filming locations",
                "Track talent mobility restrictions affecting casting and production",
            ]
        elif sector == ClientSector.ENERGY:
            actions = [
                "Brief energy sector clients on trade restrictions affecting operations",
                "Monitor oil price volatility impacts on client travel budgets",
                "Track sanctions affecting energy sector international operations",
            ]
        else:
            # Generate based on item content
            if items:
                item_content = " ".join([i.processed_content.lower() for i in items])
                if "sanction" in item_content or "restriction" in item_content:
                    actions.append("Review client operations for exposure to trade restrictions")
                if "regulation" in item_content or "compliance" in item_content:
                    actions.append("Update compliance protocols and brief affected clients")
                if "market" in item_content or "economic" in item_content:
                    actions.append("Assess market conditions impact on client travel demand")

        return actions[:3]

    def save_report(self, doc: Document, filename: Optional[str] = None) -> str:
        """Save the report to a file"""
        if not filename:
            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d")
            month = datetime.now().strftime("%B%Y")
            filename = f"Solairus_Intelligence_Report_{month}_{timestamp}.docx"

        # Get environment-appropriate output directory
        output_dir = get_output_dir()

        filepath = output_dir / filename
        doc.save(str(filepath))

        return str(filepath)


def test_document_generator():
    """Test the document generator with sample data"""
    from intelligence_processor import ClientSector, IntelligenceItem, SectorIntelligence

    # Create sample intelligence items
    sample_items = [
        IntelligenceItem(
            raw_content="Test content",
            processed_content="The Federal Reserve raised interest rates by 25 basis points, signaling continued monetary tightening.",
            category="economic_indicators",
            relevance_score=0.8,
            so_what_statement="Aircraft financing costs will increase, affecting purchase decisions and lease rates.",
            affected_sectors=[ClientSector.FINANCE],
            action_items=["Review financing arrangements", "Update cost projections"],
            confidence=0.9,
        ),
        IntelligenceItem(
            raw_content="Test content",
            processed_content="New EU regulations on sustainable aviation fuel mandate 5% SAF blend by 2025.",
            category="regulation",
            relevance_score=0.9,
            so_what_statement="Operators must plan for increased fuel costs and ensure SAF availability at European destinations.",
            affected_sectors=[ClientSector.GENERAL],
            action_items=["Establish SAF supply contracts", "Update fuel cost models"],
            confidence=0.85,
        ),
    ]

    # Create sample sector intelligence
    sector_intel = {
        ClientSector.TECHNOLOGY: SectorIntelligence(
            sector=ClientSector.TECHNOLOGY,
            items=sample_items,
            summary="Technology sector facing increased scrutiny on data sovereignty.",
            key_risks=["Export control restrictions"],
            key_opportunities=["Growing demand for tech executive travel"],
        )
    }

    # Generate document
    generator = DocumentGenerator()
    doc = generator.create_report(sample_items, sector_intel)
    filepath = generator.save_report(doc, "test_report.docx")

    print(f"Test report generated: {filepath}")
    return filepath


if __name__ == "__main__":
    test_document_generator()
