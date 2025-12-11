"""
Section builders for Solairus Intelligence Reports.

Each class is responsible for building a specific section of the report.
"""

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from solairus_intelligence.core.document.content import ContentExtractor
from solairus_intelligence.core.document.styles import SPACING, ErgoStyles
from solairus_intelligence.core.processor import (
    ClientSector,
    IntelligenceItem,
    SectorIntelligence,
)


class HeaderBuilder:
    """Builds document headers"""

    def __init__(self, styles: ErgoStyles, logo_path: Optional[str] = None):
        self.styles = styles
        self.logo_path = logo_path

    def add_header(self, doc: Document, title: str, subtitle: str) -> None:
        """Add branded header to document"""
        # Add logo if available
        if self.logo_path:
            try:
                from pathlib import Path

                if Path(self.logo_path).exists():
                    doc.add_picture(str(self.logo_path), width=Inches(1.5))
            except (FileNotFoundError, ValueError, OSError) as e:
                logging.debug(f"Could not add logo: {e}")  # Continue without logo

        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = self.styles.get_color("primary_blue")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = self.styles.get_color("dark_gray")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


class ExecutiveSummaryBuilder:
    """Builds the Executive Summary section"""

    def __init__(
        self,
        styles: ErgoStyles,
        content_extractor: ContentExtractor,
        ai_generator: Optional[Any] = None,
    ):
        self.styles = styles
        self.content_extractor = content_extractor
        self.ai_generator = ai_generator

    def add_executive_summary(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add executive summary section"""
        self._add_section_heading(doc, "Executive Summary")

        if not items:
            doc.add_paragraph("No intelligence items available for analysis.")
            return

        # Extract insights
        insights = self.content_extractor.extract_analytical_insights(items)

        # Add bottom line
        self._add_bottom_line(doc, insights.get("bottom_line", []))

        # Add key findings
        self._add_key_findings(doc, insights.get("key_findings", []))

        # Add watch factors
        self._add_watch_factors(doc, insights.get("watch_factors", []))

    def _add_section_heading(self, doc: Document, heading: str) -> None:
        """Add a section heading"""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")
        para.space_before = Pt(SPACING["section_before"])
        para.space_after = Pt(SPACING["section_after"])

    def _add_bottom_line(self, doc: Document, statements: List[str]) -> None:
        """Add bottom line up front section"""
        if not statements:
            return

        para = doc.add_paragraph()
        run = para.add_run("Bottom Line Up Front: ")
        run.font.bold = True
        run.font.size = Pt(10)

        # Add first statement inline
        if statements:
            text_run = para.add_run(statements[0])
            text_run.font.size = Pt(10)

    def _add_key_findings(self, doc: Document, findings: List[str]) -> None:
        """Add key findings section"""
        if not findings:
            return

        para = doc.add_paragraph()
        run = para.add_run("Key Findings")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.styles.get_color("secondary_blue")

        for finding in findings[:5]:
            header, desc, bullets = self.content_extractor.parse_key_finding(finding)

            # Add finding header
            finding_para = doc.add_paragraph()
            header_run = finding_para.add_run(f"• {header}: ")
            header_run.font.bold = True
            header_run.font.size = Pt(10)

            # Add description
            desc_run = finding_para.add_run(desc)
            desc_run.font.size = Pt(10)

    def _add_watch_factors(self, doc: Document, factors: List[str]) -> None:
        """Add watch factors section"""
        if not factors:
            return

        para = doc.add_paragraph()
        run = para.add_run("Watch Factors")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.styles.get_color("orange")

        for factor in factors[:3]:
            title, desc, bullets = self.content_extractor.parse_watch_factor(factor)

            factor_para = doc.add_paragraph()
            title_run = factor_para.add_run(f"⚠ {title}: ")
            title_run.font.bold = True
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = self.styles.get_color("deep_orange")

            desc_run = factor_para.add_run(desc)
            desc_run.font.size = Pt(10)


class EconomicIndicatorsBuilder:
    """Builds the Economic Indicators table"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_economic_indicators_table(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add economic indicators table"""
        # Filter for economic items
        econ_items = [
            item
            for item in items
            if item.source_type == "fred" or "economic" in item.category.lower()
        ]

        if not econ_items:
            return

        # Add section heading
        para = doc.add_paragraph()
        run = para.add_run("Economic Indicators")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")

        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        headers = ["Indicator", "Value", "Trend", "Impact"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

        # Data rows
        for item in econ_items[:6]:
            row = table.add_row()
            row.cells[0].text = self.content_extractor.extract_indicator_name(item)
            row.cells[1].text = self.content_extractor.extract_value(item)
            row.cells[2].text = self.content_extractor.determine_trend(item)
            row.cells[3].text = self.content_extractor.generate_economic_impact(item)

            # Style cells
            for cell in row.cells:
                if cell.paragraphs:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)


class RegionalAssessmentBuilder:
    """Builds the Regional Assessment section"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_regional_assessment(self, doc: Document, items: List[IntelligenceItem]) -> None:
        """Add regional assessment section"""
        # Group by region
        regions: Dict[str, List[IntelligenceItem]] = {}

        for item in items:
            region = self._detect_region(item)
            if region not in regions:
                regions[region] = []
            regions[region].append(item)

        if not regions:
            return

        # Add section heading
        para = doc.add_paragraph()
        run = para.add_run("Regional Assessment")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")

        # Add each region
        for region, region_items in regions.items():
            self._add_region_section(doc, region, region_items)

    def _detect_region(self, item: IntelligenceItem) -> str:
        """Detect region from item content"""
        content = (item.processed_content + " " + item.raw_content).lower()

        region_keywords = {
            "Europe": ["europe", "eu", "uk", "germany", "france", "italy"],
            "Asia-Pacific": ["asia", "china", "japan", "korea", "india", "pacific"],
            "Middle East": ["middle east", "saudi", "uae", "israel", "iran"],
            "Americas": ["america", "us", "usa", "canada", "mexico", "brazil"],
            "Africa": ["africa", "nigeria", "south africa", "egypt"],
        }

        for region, keywords in region_keywords.items():
            if any(kw in content for kw in keywords):
                return region

        return "Global"

    def _add_region_section(
        self, doc: Document, region: str, items: List[IntelligenceItem]
    ) -> None:
        """Add a region subsection"""
        # Region heading
        para = doc.add_paragraph()
        run = para.add_run(region)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("secondary_blue")

        # Summary of top items
        top_items = sorted(items, key=lambda x: x.relevance_score, reverse=True)[:3]

        for item in top_items:
            item_para = doc.add_paragraph()
            bullet_run = item_para.add_run("• ")
            bullet_run.font.size = Pt(10)

            if item.so_what_statement:
                text_run = item_para.add_run(item.so_what_statement)
            else:
                text_run = item_para.add_run(item.processed_content[:150] + "...")
            text_run.font.size = Pt(10)


class SectorSectionBuilder:
    """Builds sector-specific sections"""

    def __init__(self, styles: ErgoStyles, content_extractor: ContentExtractor):
        self.styles = styles
        self.content_extractor = content_extractor

    def add_sector_section(
        self, doc: Document, sector: ClientSector, intelligence: SectorIntelligence
    ) -> None:
        """Add a sector section to the document"""
        # Sector heading
        para = doc.add_paragraph()
        run = para.add_run(f"{sector.value.title()} Sector")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.styles.get_color("primary_blue")

        # Summary
        if intelligence.summary:
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(intelligence.summary)
            summary_run.font.size = Pt(10)
            summary_run.font.italic = True

        # Key items
        for item in intelligence.items[:5]:
            self._add_sector_item(doc, item, sector)

    def _add_sector_item(self, doc: Document, item: IntelligenceItem, sector: ClientSector) -> None:
        """Add a single sector item"""
        item_para = doc.add_paragraph()

        # Bullet
        bullet_run = item_para.add_run("• ")
        bullet_run.font.size = Pt(10)

        # Content
        content = item.so_what_statement or item.processed_content[:200]
        content_run = item_para.add_run(content)
        content_run.font.size = Pt(10)

        # Action items if available
        if item.action_items:
            for action in item.action_items[:2]:
                action_para = doc.add_paragraph()
                action_para.paragraph_format.left_indent = Inches(0.25)
                action_run = action_para.add_run(f"→ {action}")
                action_run.font.size = Pt(9)
                action_run.font.color.rgb = self.styles.get_color("dark_gray")
